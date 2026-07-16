from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_steadfast_listing_workflow as wf

ORIGINAL_TABLE_WIDTHS = wf.table_widths


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SteadFast_Database_Design_v0.1.md"
OUTPUT = ROOT / "docs" / "SteadFast_Database_Design_v0.1.docx"
TEMPLATE = ROOT / "docs" / "SteadFast_MVP_Product_Requirements_v0.1.docx"
ASSET_DIR = ROOT / "docs" / "assets"
ERD_IMAGE = ASSET_DIR / "SteadFast_Database_ERD_v0.1.png"


def load_font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def build_erd():
    ASSET_DIR.mkdir(exist_ok=True)
    width, height = 1900, 1280
    image = Image.new("RGB", (width, height), "#F7FAFC")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, True)
    subtitle_font = load_font(23)
    box_font = load_font(23, True)
    small_font = load_font(17)

    draw.text((70, 40), "SteadFast Core Entity Relationship Map", font=title_font, fill="#17365D")
    draw.text((70, 95), "Brokerage-rooted identity, listing workflow, publication, billing, integration, and audit domains", font=subtitle_font, fill="#667085")

    colors = {
        "identity": ("#DCEBFA", "#1F4E78"),
        "brokerage": ("#E5F3E8", "#287D3C"),
        "listing": ("#FFF0D8", "#8A5A00"),
        "service": ("#ECE7F6", "#5B3F8C"),
        "platform": ("#E9EDF2", "#344054"),
    }
    nodes = {
        "auth_users": (90, 190, "auth.users", "Supabase identity", "identity"),
        "people": (520, 190, "people", "Durable person", "identity"),
        "memberships": (950, 190, "brokerage_memberships", "One active brokerage", "brokerage"),
        "brokerages": (1380, 190, "brokerages", "Own listings", "brokerage"),
        "roles": (90, 410, "roles + permissions", "Delegated authority", "identity"),
        "subscriptions": (520, 410, "subscriptions + seats", "One paid professional seat", "service"),
        "listings": (950, 410, "listings", "Lifecycle + approved pointers", "listing"),
        "properties": (1380, 410, "properties + addresses", "PostGIS location", "listing"),
        "sites": (90, 630, "sites + domains", "Agent/brokerage web", "service"),
        "assignments": (520, 630, "listing_assignments", "Active representative", "brokerage"),
        "versions": (950, 630, "listing_versions", "Immutable snapshots", "listing"),
        "reviews": (1380, 630, "listing_reviews", "Approve/return/reject", "listing"),
        "media": (90, 850, "media_assets", "Private source/public derivative", "service"),
        "shares": (520, 850, "listing_shares", "Display permission only", "brokerage"),
        "snapshots": (950, 850, "public_listing_snapshots", "Sanitized public search", "listing"),
        "inquiries": (1380, 850, "inquiries + recipients", "Selected/secondary agent", "service"),
        "channels": (90, 1070, "channels + distributions", "External delivery", "platform"),
        "billing": (520, 1070, "invoices + payments", "Append-only finance", "service"),
        "notifications": (950, 1070, "notifications + outbox", "Reliable delivery", "platform"),
        "audit": (1380, 1070, "audit_events", "Immutable history", "platform"),
    }
    box_w, box_h = 330, 120

    centers = {key: (x + box_w // 2, y + box_h // 2) for key, (x, y, _, _, _) in nodes.items()}
    edges = [
        ("auth_users", "people"), ("people", "memberships"), ("brokerages", "memberships"),
        ("memberships", "roles"), ("people", "subscriptions"), ("brokerages", "listings"),
        ("properties", "listings"), ("memberships", "assignments"), ("listings", "assignments"),
        ("listings", "versions"), ("versions", "reviews"), ("versions", "media"),
        ("listings", "shares"), ("memberships", "shares"), ("listings", "snapshots"),
        ("listings", "inquiries"), ("people", "inquiries"), ("subscriptions", "billing"),
        ("listings", "channels"), ("versions", "notifications"), ("listings", "audit"),
        ("people", "audit"), ("sites", "snapshots"),
    ]
    for start, end in edges:
        x1, y1 = centers[start]
        x2, y2 = centers[end]
        draw.line((x1, y1, x2, y2), fill="#AAB4C0", width=4)
        dx, dy = x2 - x1, y2 - y1
        mag = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / mag, dy / mag
        px, py = x2 - ux * (box_w * 0.45), y2 - uy * (box_h * 0.45)
        left = (px - ux * 18 - uy * 9, py - uy * 18 + ux * 9)
        right = (px - ux * 18 + uy * 9, py - uy * 18 - ux * 9)
        draw.polygon([(px, py), left, right], fill="#7E8A97")

    for key, (x, y, title, subtitle, domain) in nodes.items():
        fill, outline = colors[domain]
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=fill, outline=outline, width=4)
        title_box = draw.textbbox((0, 0), title, font=box_font)
        title_w = title_box[2] - title_box[0]
        if title_w > box_w - 24:
            title_font_used = load_font(19, True)
        else:
            title_font_used = box_font
        draw.text((x + 18, y + 24), title, font=title_font_used, fill=outline)
        draw.text((x + 18, y + 70), subtitle, font=small_font, fill="#344054")

    legend_y = height - 48
    draw.text((90, legend_y), "Arrows show principal foreign-key or event-flow relationships; full cardinality and constraints are defined in the document.", font=small_font, fill="#667085")
    image.save(ERD_IMAGE, "PNG", optimize=True)


def db_table_widths(count):
    choices = {
        2: [2300, 7060],
        3: [1800, 3000, 4560],
        4: [1600, 2500, 1900, 3360],
        5: [1400, 1500, 1500, 2200, 2760],
        6: [1250, 1350, 1350, 2100, 1450, 1860],
    }
    return choices.get(count, ORIGINAL_TABLE_WIDTHS(count))


def configure_header(doc):
    wf.configure_styles(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    wf.set_font(header.add_run("STEADFAST  |  DATABASE DESIGN"), size=8.5, bold=True, color=wf.MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    wf.set_font(footer.add_run("Target Schema  |  Version 0.1  |  Page "), size=9, color=wf.MID_GRAY)
    wf.add_page_number(footer)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    wf.set_font(p.add_run("DATA ARCHITECTURE SPECIFICATION"), size=10, bold=True, color=wf.BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    wf.set_font(p.add_run("SteadFast Database Design"), size=28, bold=True, color=wf.DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    wf.set_font(p.add_run("Brokerage-isolated PostgreSQL and Supabase target model"), size=14, color=wf.MID_GRAY)
    for label, value in (
        ("Document owner", "SteadFast Product and Engineering"),
        ("Status", "Target schema; no database changes performed"),
        ("Version", "0.1"),
        ("Date", "16 July 2026"),
        ("Primary runtime", "PostgreSQL on Supabase with PostGIS"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        wf.set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=wf.DARK_BLUE)
        wf.set_font(p.add_run(value), size=10.5, color=wf.DARK)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    wf.set_table_geometry(table, [9360])
    wf.set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    wf.set_cell_shading(cell, wf.PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    wf.set_font(p.add_run("Security baseline: "), size=11, bold=True, color=wf.BLUE)
    wf.set_font(p.add_run("Anonymous users query sanitized approved projections. Raw listings, exact addresses, drafts, reviews, billing, support, and audit data remain protected by explicit grants and Row Level Security."), size=11, color=wf.DARK)
    doc.add_paragraph()


def add_erd_picture(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline = run.add_picture(str(ERD_IMAGE), width=Inches(6.5))
    doc_pr = inline._inline.docPr
    doc_pr.set("name", "SteadFast core entity relationship map")
    doc_pr.set("descr", "Core SteadFast database relationships across identity, brokerage, listings, publication, inquiries, billing, integrations, and audit.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    wf.set_font(caption.add_run("Figure 1. Core entity relationship map. Detailed fields and constraints follow."), size=9, italic=True, color=wf.MID_GRAY)


def build():
    build_erd()
    wf.table_widths = db_table_widths
    doc = Document(TEMPLATE)
    wf.clear_body(doc)
    configure_header(doc)
    add_masthead(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("**Version:") or line.startswith("**Prepared:") or line.startswith("**Status:") or line.startswith("**Platform:") or line.startswith("**Applies to:"):
            i += 1
            continue
        if line == "```mermaid":
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                i += 1
            i += 1
            add_erd_picture(doc)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = wf.parse_table(lines, i)
            wf.add_table(doc, rows)
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            wf.add_inline(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            wf.add_inline(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        wf.add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = "SteadFast Database Design v0.1"
    props.subject = "PostgreSQL and Supabase target schema, ERD, RLS, geography, storage, integration, billing, audit, and migration design"
    props.author = "SteadFast"
    props.keywords = "SteadFast, PostgreSQL, Supabase, database, ERD, RLS, PostGIS"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
