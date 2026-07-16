from pathlib import Path
import runpy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)
PROPOSAL_PATH = DOCS / "SteadFast_Business_Proposal_v0.1.docx"
OUTPUT = DOCS / "SteadFast_Business_Plan_v0.1.docx"

# Load the established SteadFast proposal design and helper functions.
ns = runpy.run_path(str(ROOT / "tools" / "build_steadfast_business_proposal.py"))
doc = Document(str(PROPOSAL_PATH))
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

BLUE = ns["BLUE"]
DARK_BLUE = ns["DARK_BLUE"]
DARK = ns["DARK"]
MID_GRAY = ns["MID_GRAY"]
LIGHT_BLUE = ns["LIGHT_BLUE"]
set_font = ns["set_font"]
add_heading = ns["add_heading"]
add_body = ns["add_body"]
add_bullets = ns["add_bullets"]
add_numbered = ns["add_numbered"]
add_callout = ns["add_callout"]
add_table = ns["add_table"]
add_page_number = ns["add_page_number"]

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Narrative proposal preset: Calibri 11, justified, 8 pt after, 1.333 spacing.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("STEADFAST  |  BUSINESS PLAN"), size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Confidential Planning Draft  |  Version 0.1  |  Page "), size=9, color=MID_GRAY)
add_page_number(fp)


def centered(text, size=11, bold=False, color=DARK, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_lead(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.333
    set_font(p.add_run(text), size=12, bold=True, color=DARK_BLUE)
    return p


# Editorial business-plan cover.
centered("STEADFAST", size=12, bold=True, color=BLUE, after=24)
centered("Business Plan", size=30, bold=True, color=DARK_BLUE, after=7)
centered("Jamaica-first real-estate listing and marketing SaaS", size=16, bold=True, color=BLUE, after=12)
centered(
    "A broker-controlled platform for listing operations, agent collaboration, public property discovery, branded websites, and authorized international distribution",
    size=11.5,
    color=MID_GRAY,
    after=28,
)
add_table(doc, ["Plan profile", "Details"], [
    ("Company", "SteadFast"),
    ("Launch market", "Jamaica"),
    ("Expansion direction", "Belize, followed by selected Caribbean and Central American markets"),
    ("Business model", "Subscription software for agents, broker staff, and brokerages; free consumer access"),
    ("Status", "Planning baseline - Version 0.1"),
    ("Date", "16 July 2026"),
], [2400, 6960], header_fill=LIGHT_BLUE, font_size=10)
add_callout(doc, "Planning note", "All financial figures are illustrative management assumptions, not forecasts or investment guarantees. They must be replaced with pilot evidence, confirmed costs, and signed customer commitments.")

add_heading(doc, "1. Executive summary", 1)
add_lead("SteadFast will be a cloud operating platform for real-estate brokers and agents, beginning in Jamaica.")
add_body(doc, "The business will solve a practical industry problem: professionals need a simple way to create, approve, market, share, and distribute property listings without relying on disconnected tools or difficult legacy workflows. SteadFast combines broker-controlled listing operations with agent websites, brokerage websites, map-based public search, inquiries, subscription management, audit history, and integration-ready property data.")
add_body(doc, "The company will launch with a paid pilot among interested Jamaican brokerages and agents. Visitors and registered consumers will use the marketplace free of charge. Professional agents will pay US$50 per month. Brokerages will pay an additional company subscription based on the number of affiliated agents and non-agent staff they manage. The sponsor reports early interest from approximately 200 professionals; this is a valuable lead pool but must not be treated as paid demand until contracts and payments are secured.")
add_body(doc, "The initial objective is to prove adoption, listing quality, approval efficiency, customer inquiries, retention, and willingness to pay in Jamaica. Belize is the preferred second market. Additional Caribbean and Central American markets will be considered individually after local partnerships, operating requirements, and authorized distribution arrangements are confirmed.")

add_heading(doc, "2. Company overview", 1)
add_heading(doc, "2.1 Mission", 2)
add_body(doc, "To give Caribbean real-estate professionals a trustworthy, easy-to-use platform that improves listing quality, brokerage control, agent visibility, collaboration, and access to buyers.")
add_heading(doc, "2.2 Vision", 2)
add_body(doc, "To become the leading broker-centered listing and property marketing infrastructure for Jamaica and selected regional markets, connecting accurate local inventory with domestic and international buyers.")
add_heading(doc, "2.3 Strategic principles", 2)
add_bullets(doc, [
    "Brokerage accountability: brokerages own and approve listings while agents provide the required public representation.",
    "Ease of use: a novice agent should complete core work without formal training.",
    "Accurate, controlled data: publication, sharing, changes, and removals must be authorized and traceable.",
    "Professional opportunity: every paid professional receives useful tools and a credible public presence.",
    "Portability: technology, data, documentation, and deployments must not depend permanently on one hosting provider.",
    "Regional readiness: country, currency, language, location, policy, and integration differences must be configurable.",
])

add_heading(doc, "3. Market problem and opportunity", 1)
add_body(doc, "Real-estate listing platforms often emphasize public advertising while under-serving the operational relationship between brokerages and agents. The Jamaican regulatory model makes that relationship especially important: the Real Estate Board distinguishes dealers from salesmen and describes salesmen as working under a dealer's direction, control, or management. SteadFast's broker-rooted structure fits that operating reality.")
add_bullets(doc, [
    "Agents need fast listing tools, professional websites, shared inventory, customer inquiries, and mobile access.",
    "Brokers need approvals, ownership, agent and staff management, reporting, reassignment, and complete history.",
    "Consumers need accurate search, maps, strong media, and a clear path to an accountable local professional.",
    "Brokerages need a consistent online presence without maintaining several separate websites and databases.",
    "International distribution requires structured data, content rights, update controls, and operational monitoring.",
    "Regional markets present an opportunity for a configurable platform built around local professional relationships rather than a single-country portal.",
])
add_callout(doc, "Initial market signal", "Approximately 200 professionals are reported to be interested. The paid pilot must determine how many are qualified, active, willing to migrate listings, and willing to subscribe at the agreed price.")

add_heading(doc, "4. Target customers", 1)
add_table(doc, ["Segment", "Need", "SteadFast response"], [
    ("Individual agents", "Listings, visibility, collaboration, leads, and a professional website", "US$50 monthly professional subscription with agent workspace and subdomain website."),
    ("Small brokerages", "Control and branding without enterprise complexity", "Broker Core for up to 20 agents and 5 non-agent staff."),
    ("Growing brokerages", "Larger teams, stronger reporting, and support", "Broker Growth for up to 75 agents and 15 non-agent staff."),
    ("Large networks", "Scale, migration, integration, and service commitments", "Custom Enterprise plan above Growth limits."),
    ("Property consumers", "Free, accurate, easy property discovery and professional contact", "Public marketplace and optional free registered account."),
    ("Distribution partners", "Structured, authorized, monitored listing data", "RESO-aligned integration layer and channel-specific adapters."),
], [1800, 3100, 4460], font_size=9.1)

add_heading(doc, "5. Product and services", 1)
add_heading(doc, "5.1 Core platform", 2)
add_bullets(doc, [
    "Residential sale, long-term rental, commercial, land, and development listings; vacation rental operations are excluded initially.",
    "Drafts, broker approval, change comparison, return comments, price and status controls, removal, reassignment, and audit history.",
    "Public, professional-network, and private visibility settings governed by listing status and permissions.",
    "Agent-to-agent display sharing without ownership transfer or editing authority.",
    "Agent and brokerage subdomain websites that update automatically from approved listing data.",
    "Public search, filters, property pages, clustered maps, agent selection, inquiries, favourites, and saved searches.",
    "Professional dashboards for agents, broker staff, brokers, SteadFast operations, and administrators.",
    "Subscriptions, invoices, plan limits, notifications, support records, flags, and integration monitoring.",
])
add_heading(doc, "5.2 Integration direction", 2)
add_body(doc, "SteadFast will maintain its own authoritative listing record and map data to RESO-aligned property, member, office, and media concepts where practical. The platform will expose documented interfaces and separate external adapters so that changes to one destination do not disrupt the core business workflow.")
add_body(doc, "Move operates Realtor.com, but technical compatibility does not grant distribution rights. Realtor.com International, Move-operated channels, RAJ/MLS, and future destinations must each be treated as governed commercial integrations requiring authorization, accepted data specifications, monitoring, and removal procedures.")

add_heading(doc, "6. Competitive strategy", 1)
add_table(doc, ["Alternative", "Typical strength", "SteadFast advantage"], [
    ("Legacy or local listing systems", "Existing inventory and industry familiarity", "Simpler workflows, modern interface, websites, sharing, maps, and accountable change history."),
    ("Global property portals", "Large consumer audiences", "Local brokerage operations, ownership, approvals, agent relationships, and regional configuration."),
    ("Standalone agent websites", "Personal branding", "Automatically synchronized approved inventory, brokerage connection, collaboration, and no separate content maintenance."),
    ("Generic CRM or form tools", "Flexible customer tracking", "Purpose-built property data, broker approval, distribution status, maps, and public listing lifecycle."),
], [2100, 3000, 4260], font_size=9.1)
add_heading(doc, "6.1 Defensible capabilities", 2)
add_bullets(doc, [
    "Brokerage and agent network relationships supported by a product advisory group.",
    "Accumulated listing, approval, sharing, inquiry, and distribution history.",
    "Country-specific workflows and location data that generic portals do not prioritize.",
    "Low-friction agent and brokerage websites embedded in the core subscription.",
    "Reliable integrations and removal controls developed through partner agreements.",
    "Operational documentation and a portable technical platform that can be expanded without a complete rewrite.",
])

add_heading(doc, "7. Revenue model and pricing", 1)
add_table(doc, ["Plan", "Monthly price", "Capacity and purpose"], [
    ("Consumer", "Free", "Visitor search and optional registered-user tools."),
    ("Professional Agent", "US$50", "One professional account, listing tools when authorized, agent website, collaboration, inquiries, notifications, and support."),
    ("Broker Core", "US$150", "Brokerage administration for up to 20 affiliated agents and 5 non-agent staff; agent subscriptions remain separate."),
    ("Broker Growth", "US$300", "Brokerage administration for up to 75 affiliated agents and 15 non-agent staff, stronger reporting, and priority onboarding."),
    ("Enterprise", "Custom", "More than 75 agents or 15 non-agent staff, migration, integration, negotiated service levels, and volume arrangements."),
], [1900, 1500, 5960], font_size=9.1)
add_callout(doc, "Role and billing rule", "A person pays once even when holding several roles. An agent with broker-staff permissions counts as an affiliated agent, not as an additional non-agent staff account. A brokerage may pay professional subscriptions centrally, but the underlying active seats remain identifiable.")
add_heading(doc, "7.1 Additional future revenue", 2)
add_bullets(doc, [
    "Custom domains and premium website themes.",
    "Advanced analytics and brokerage reporting.",
    "Premium onboarding, data migration, and listing-quality services.",
    "Authorized international distribution packages where partner terms allow resale or bundled pricing.",
    "Promoted inventory and featured placement with clear advertising labels.",
    "Approved data exports, CRM connections, and custom enterprise integrations.",
])

add_heading(doc, "8. Sales and go-to-market strategy", 1)
add_heading(doc, "8.1 Jamaica launch", 2)
add_numbered(doc, [
    "Qualify the reported 200 interested professionals by brokerage, role, current tools, listing volume, decision authority, and willingness to pay.",
    "Recruit five to ten founding brokerages representing different sizes and operating styles.",
    "Run structured product demonstrations and obtain written pilot commitments.",
    "Launch a controlled paid pilot with real listings, weekly feedback, onboarding support, and measurable success criteria.",
    "Convert successful pilot accounts to standard subscriptions and publish approved customer stories.",
    "Grow through broker referrals, agent referrals, industry events, educational demonstrations, and professional relationships.",
])
add_heading(doc, "8.2 Sales funnel", 2)
add_table(doc, ["Stage", "Management measure"], [
    ("Lead", "Named professional or brokerage with source and contact permission."),
    ("Qualified", "Confirmed role, brokerage relationship, listing activity, need, authority, and price discussion."),
    ("Pilot committed", "Signed pilot terms, billing method, onboarding date, and success owner."),
    ("Activated", "Account configured and first approved listing or brokerage workflow completed."),
    ("Paying", "Successful payment and recurring subscription in good standing."),
    ("Retained", "Continued active use, acceptable support burden, renewal, and referral potential."),
], [1800, 7560], font_size=9.4)

add_heading(doc, "9. Customer success and retention", 1)
add_body(doc, "SteadFast is intended to require little or no formal training, but successful adoption still needs a deliberate customer-success process. Brokerages must be configured correctly, decision-makers must see operational value, and agents must reach their first useful outcome quickly.")
add_bullets(doc, [
    "Guided brokerage setup with roles, staff permissions, branding, agent invitations, and listing rules.",
    "In-product checklists, plain-language help, examples, and short demonstration videos.",
    "Measured time to first submitted listing, first approval, first website publication, and first inquiry.",
    "Support classification for product questions, billing, defects, data issues, and integration failures.",
    "Monthly broker health reviews during the pilot and automated usage signals after launch.",
    "Cancellation and exit process that preserves required records while clearly ending public display and access.",
])

add_heading(doc, "10. Operations plan", 1)
add_heading(doc, "10.1 Initial operating functions", 2)
add_table(doc, ["Function", "Initial responsibility"], [
    ("Product leadership", "Customer discovery, priorities, acceptance decisions, pricing, and roadmap."),
    ("Engineering", "Application, database, security, integrations, testing, deployment, and technical documentation."),
    ("Customer operations", "Onboarding, account support, billing assistance, issue triage, and response tracking."),
    ("Broker partnerships", "Pilot recruitment, relationship management, feedback, expansion partners, and distribution discussions."),
    ("Finance and administration", "Subscriptions, collections, bookkeeping, taxes, vendors, budgeting, and management reporting."),
    ("Legal and compliance", "Terms, privacy, data processing, content rights, professional rules, partner contracts, and incident advice."),
], [2400, 6960], font_size=9.2)
add_heading(doc, "10.2 Service boundaries", 2)
add_bullets(doc, [
    "Brokerages approve listings and resolve listing, agent, commission, ownership, and professional disputes.",
    "SteadFast handles platform availability, access, billing support, product defects, integration failures, and communications.",
    "SteadFast may monitor flags, preserve evidence, and notify the responsible broker without replacing the broker's decision-making role.",
    "A narrowly defined legal and security emergency process must be approved before production launch.",
])

add_heading(doc, "11. Technology and delivery plan", 1)
add_body(doc, "The prototype will be developed locally and published through client-viewable preview deployments. The initial stack will use a modern web application, PostgreSQL-based data, secure authentication, file storage, and geospatial search. Vercel and Supabase are suitable for the prototype and early service, while deployment files, migrations, interfaces, and documentation will preserve the option to move to other compatible infrastructure.")
add_bullets(doc, [
    "Private Git repository owned by the SteadFast company, protected main branch, reviewed changes, and automated checks.",
    "Local, feature-preview, client-demo, and production environments with separate data and credentials.",
    "Versioned database migrations, backups, recovery tests, monitoring, and traceable deployments.",
    "Secure server-side authorization and database row-level controls to isolate brokerage data.",
    "PostGIS or equivalent geospatial capability for property coordinates, map bounds, and clustering.",
    "Technical documentation covering product rules, architecture, database, APIs, operations, security, deployment, and developer onboarding.",
])

add_heading(doc, "12. Legal, regulatory, and security approach", 1)
add_body(doc, "The Real Estate Board regulates Jamaica's private real-estate sector and licenses dealers and salesmen. The platform should verify professional registration details as part of brokerage onboarding while leaving formal licensing decisions to the regulator and brokerages. SteadFast must also obtain Jamaican legal and tax advice before production launch.")
add_bullets(doc, [
    "Comply with applicable provisions of Jamaica's Data Protection Act, 2020, including lawful processing, notices, data-subject rights, security, retention, and processor arrangements.",
    "Confirm content ownership and permission for property descriptions, photographs, documents, map locations, and external distribution.",
    "Use clear subscription terms, acceptable-use rules, privacy notices, refund terms, and broker responsibilities.",
    "Protect authentication, sessions, uploads, secrets, administrative tools, logs, and external integration credentials.",
    "Apply least privilege, multi-factor authentication for powerful roles, rate limits, monitoring, backups, and incident response.",
    "Complete a formal threat model before the external pilot and an independent security review before full production launch.",
])

add_heading(doc, "13. Partnerships and international growth", 1)
add_heading(doc, "13.1 Priority relationships", 2)
add_bullets(doc, [
    "Founding Jamaican brokerages and agent advisers.",
    "Real Estate Board and professional-association contacts for requirements and education, without implying endorsement.",
    "RAJ/MLS or applicable listing-data authorities for authorized data exchange.",
    "Realtor.com International, Move, ListHub, or the appropriate distribution program for international exposure.",
    "Payment, mapping, messaging, email, storage, monitoring, and security vendors.",
    "Belizean anchor brokerages and local professional, legal, tax, payment, and data partners before second-market launch.",
])
add_heading(doc, "13.2 Expansion gate", 2)
add_body(doc, "A new market should launch only after SteadFast confirms an anchor customer, professional and legal requirements, country locations, currency, language and terminology, tax treatment, payment methods, content and privacy rules, support coverage, and a realistic acquisition plan. International exposure is a distribution capability, not a substitute for local product-market fit.")

add_heading(doc, "14. Milestones", 1)
add_table(doc, ["Period", "Primary milestones"], [
    ("0-3 months", "Confirm pilot partners, scope, brand, repository, prototype, account hierarchy, security baseline, and documented product decisions."),
    ("4-6 months", "Deliver listing drafts, approval workflow, brokerage ownership, media, audit history, and internal pilot testing."),
    ("7-9 months", "Deliver public search, maps, property pages, inquiries, agent and brokerage websites, sharing, and notifications."),
    ("10-12 months", "Run paid Jamaica pilot, introduce subscriptions and support operations, measure adoption, and prepare launch decision."),
    ("Year 2", "Expand Jamaican customers, improve reporting and integrations, pursue authorized international distribution, and prepare Belize pilot."),
    ("Year 3", "Scale Jamaica and Belize, evaluate additional markets, strengthen enterprise services, and optimize profitability."),
], [1800, 7560], font_size=9.2)

add_heading(doc, "15. Financial plan", 1)
add_heading(doc, "15.1 Core financial assumptions", 2)
add_table(doc, ["Assumption", "Planning value"], [
    ("Professional agent subscription", "US$50 per month per active paid professional seat."),
    ("Broker Core", "US$150 per month for up to 20 affiliated agents and 5 non-agent staff."),
    ("Broker Growth", "US$300 per month for up to 75 affiliated agents and 15 non-agent staff."),
    ("Enterprise", "Custom price for customers above Growth limits or requiring negotiated services."),
    ("Consumer access", "Free; no consumer revenue assumed in the base plan."),
    ("Professional interest", "Approximately 200 reported leads; conversion and retention are unverified."),
    ("Annual prepayment", "Not included in the base model; may be introduced after pilot retention is understood."),
], [3000, 6360], font_size=9.3)

add_heading(doc, "15.2 Illustrative end-of-year recurring revenue", 2)
add_table(doc, ["Metric", "Year 1", "Year 2", "Year 3"], [
    ("Paid professional seats", "150", "400", "900"),
    ("Broker Core accounts", "6", "15", "30"),
    ("Broker Growth accounts", "1", "4", "10"),
    ("Professional seat ARR", "US$90,000", "US$240,000", "US$540,000"),
    ("Broker plan ARR", "US$14,400", "US$41,400", "US$90,000"),
    ("Total ending ARR", "US$104,400", "US$281,400", "US$630,000"),
], [3000, 2120, 2120, 2120], font_size=9.1)
add_body(doc, "Ending annual recurring revenue measures the annualized value of active subscriptions at year-end; it is not the same as revenue recognized during the year. Because subscriptions are expected to ramp gradually, recognized first-year revenue would be lower than ending ARR. A cash-flow model should be completed once development staffing, vendor, legal, insurance, marketing, payment, and support costs are quoted.")

add_heading(doc, "15.3 Management targets", 2)
add_table(doc, ["Measure", "Initial target"], [
    ("Pilot-to-paid conversion", "At least 60% of activated pilot professionals."),
    ("Monthly professional churn", "Below 2.5% after the first six months of paid service."),
    ("Gross margin", "Long-term target of 75%-85%, excluding early product-development investment."),
    ("Customer acquisition payback", "Below six months of subscription gross profit."),
    ("First-value time", "First submitted listing or configured brokerage workflow within one business day of onboarding."),
    ("Support load", "Track contacts per active professional and resolve repeated causes through product improvements."),
], [3000, 6360], font_size=9.3)
add_callout(doc, "Financial discipline", "Do not hire or commit to regional expansion based only on ending-ARR projections. Use collected cash, retention, active usage, support demand, and signed pipeline to release each spending stage.")

add_heading(doc, "16. Funding and cost strategy", 1)
add_body(doc, "The prototype can begin with low-cost development and free or entry-level hosting tiers, but the production business must budget for reliable infrastructure, backups, monitoring, email and messaging, maps, storage, payment processing, security, legal advice, accounting, customer support, and engineering. Free plans are a prototype tool, not a production operating model.")
add_bullets(doc, [
    "Stage 1: founder-funded discovery and prototype with tightly controlled scope.",
    "Stage 2: paid pilot revenue and a defined launch budget for production reliability and support.",
    "Stage 3: reinvest recurring revenue into sales, customer success, integrations, and measured regional growth.",
    "External funding, if pursued, should be tied to verified retention, paid brokerage adoption, and a credible expansion pipeline.",
])

add_heading(doc, "17. Key risks and responses", 1)
add_table(doc, ["Risk", "Response"], [
    ("Interest does not convert to paid use", "Obtain signed paid pilots, measure activation and retention, and delay large fixed costs."),
    ("Agents resist US$50 monthly", "Test value messaging and usage; consider limited founding discounts without permanently weakening pricing."),
    ("Broker approval slows publication", "Provide clear queues, change comparisons, notifications, delegates, and turnaround reporting."),
    ("International feed is unavailable", "Maintain a strong domestic product and do not promise distribution before written approval."),
    ("Incorrect or unauthorized listings", "Broker approval, content warranties, audit history, flags, removal process, and clear accountability."),
    ("Security or privacy incident", "Secure-by-default architecture, monitoring, backups, incident procedures, threat modeling, and independent review."),
    ("Support costs grow too quickly", "Self-service onboarding, product analytics, recurring-issue fixes, plan limits, and documented service boundaries."),
    ("Expansion distracts from Jamaica", "Require measurable Jamaica retention and a market-entry gate before launching another country."),
], [3000, 6360], font_size=9.1)

add_heading(doc, "18. Performance dashboard", 1)
add_bullets(doc, [
    "Qualified brokers, qualified professionals, pilot commitments, activated users, paid users, and conversion rates.",
    "Monthly recurring revenue, ending ARR, collections, failed payments, refunds, discounts, and revenue concentration.",
    "Active brokerages, active agents, listings created, approval time, public listings, shared displays, inquiries, and inquiry response.",
    "First-value time, weekly and monthly active professionals, retention, churn, reactivation, and referrals.",
    "Support contacts, response time, resolution time, repeat issues, incidents, and customer satisfaction.",
    "Infrastructure cost per active professional, storage per listing, map and messaging usage, and gross-margin trend.",
    "Feed acceptance, validation errors, update delay, removal delay, and external distribution availability when integrations launch.",
])

add_heading(doc, "19. Immediate action plan", 1)
add_numbered(doc, [
    "Confirm the legal company structure, ownership, decision authority, and banking approach for SteadFast.",
    "Turn the reported 200 interested professionals into a named, permissioned, qualified lead list organized by brokerage.",
    "Recruit the founding broker advisory group and obtain written pilot participation terms.",
    "Approve the launch scope, role rules, pricing, pilot success measures, and customer agreements.",
    "Create the private company Git repository, prototype environments, technical documentation structure, and security baseline.",
    "Build and demonstrate the listing and broker-approval workflow before expanding into secondary features.",
    "Obtain formal legal, tax, privacy, payment, insurance, and international-distribution advice.",
    "Review pilot evidence and collected revenue before approving public launch, major hiring, or Belize expansion.",
])

add_heading(doc, "20. Sources and planning basis", 1)
add_bullets(doc, [
    "Real Estate Board of Jamaica regulatory role and dealer/salesman framework: https://reb.gov.jm/about-us/our-story/ and https://reb.gov.jm/frequently-asked-questions/",
    "Jamaica Data Protection Act, 2020: https://laws.moj.gov.jm/library/act-of-parliament/7-2020-the-data-protection-act",
    "Move and Realtor.com operating relationship: https://www.move.com/",
    "RESO Web API and data-authorization overview: https://www.reso.org/reso-web-api/",
    "SteadFast Product and Technical Plan, Version 0.1, and SteadFast Business Proposal, Version 0.1.",
])
add_body(doc, "This plan should be reviewed quarterly during development and monthly during the paid pilot. Assumptions should be replaced with actual pipeline, product usage, revenue, retention, cost, and customer-support evidence as it becomes available.")

settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "SteadFast Business Plan"
doc.core_properties.subject = "Jamaica-first real-estate SaaS business plan"
doc.core_properties.author = "SteadFast"
doc.core_properties.keywords = "SteadFast, business plan, Jamaica, real estate, SaaS, brokers, agents"
doc.save(str(OUTPUT))
print(OUTPUT)
