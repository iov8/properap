from pathlib import Path
import runpy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)
TEMPLATE = DOCS / "SteadFast_Business_Plan_v0.1.docx"
OUTPUT = DOCS / "SteadFast_MVP_Product_Requirements_v0.1.docx"

ns = runpy.run_path(str(ROOT / "tools" / "build_steadfast_business_plan.py"))
doc = Document(str(TEMPLATE))
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

BLUE = ns["BLUE"]
DARK_BLUE = ns["DARK_BLUE"]
DARK = ns["DARK"]
MID_GRAY = ns["MID_GRAY"]
LIGHT_BLUE = ns["LIGHT_BLUE"]
set_font = ns["set_font"]
add_heading = ns["add_heading"]
add_body = ns["add_body"]
add_bullets = ns["add_bullets"]
add_numbered = ns["add_numbered"]
add_callout = ns["add_callout"]
add_table = ns["add_table"]
add_page_number = ns["add_page_number"]

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Standard business brief preset.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("STEADFAST  |  MVP PRODUCT REQUIREMENTS"), size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Planning Baseline  |  Version 0.1  |  Page "), size=9, color=MID_GRAY)
add_page_number(fp)


def add_req(req_id, title, requirement, acceptance):
    add_heading(doc, f"{req_id} - {title}", 3)
    add_body(doc, requirement, bold_lead="Requirement: ")
    add_body(doc, acceptance, bold_lead="Acceptance: ")


def add_kv(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
    set_font(p.add_run(value), color=DARK)


# Memo masthead.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(3)
set_font(p.add_run("PRODUCT REQUIREMENTS DOCUMENT"), size=10, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("SteadFast MVP"), size=29, bold=True, color=DARK_BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
set_font(p.add_run("Jamaica pilot release for broker-controlled property listing and marketing"), size=14, color=MID_GRAY)

add_kv("Document owner", "SteadFast Product Management")
add_kv("Status", "Planning baseline for stakeholder review")
add_kv("Version", "0.1")
add_kv("Date", "16 July 2026")
add_kv("Primary audience", "Product, engineering, design, testing, operations, and future IT staff")
doc.add_paragraph()
add_callout(doc, "MVP objective", "Prove that Jamaican agents and brokerages can create, approve, publish, share, discover, and manage listings in a simple paid platform that requires little or no training.")

add_heading(doc, "1. Purpose", 1)
add_body(doc, "This Product Requirements Document defines the first buildable SteadFast release for a controlled Jamaica pilot. It translates the approved product and business direction into testable requirements. The separate permissions matrix, workflow specification, database design, architecture, security plan, and development backlog will expand the implementation detail without changing the product rules recorded here.")

add_heading(doc, "2. Product summary", 1)
add_body(doc, "SteadFast is a multi-broker cloud platform for property listing operations and marketing. The brokerage is the professional root: it owns listings, manages agents and staff, approves public content, and receives a branded brokerage website. Agents receive listing tools and personal websites. Consumers receive free public search, maps, property pages, and agent contact options.")
add_bullets(doc, [
    "Launch market: Jamaica.",
    "Initial property scope: residential sale, long-term rental, commercial, land, and developments.",
    "Excluded property operation: vacation and short-term rental management.",
    "Primary professional model: one agent belongs to one brokerage at a time.",
    "Primary data model: brokerage-owned listings with an active agent representative required for public display.",
    "International direction: integration-ready records; live distribution requires separate authorization and is not an MVP dependency.",
])

add_heading(doc, "3. MVP goals", 1)
add_bullets(doc, [
    "Allow a qualified professional to join a brokerage and reach productive use quickly.",
    "Allow an agent to create and submit a complete listing without training.",
    "Allow a broker or authorized staff member to review, compare, approve, return, reject, and trace listing work.",
    "Prevent unapproved, inactive, or unassigned listings from appearing publicly.",
    "Provide visitors with fast list and map search and a direct path to an agent.",
    "Provide every active agent and brokerage with an automatically maintained public website.",
    "Support display-only sharing between agents with correct contacts and notifications.",
    "Enforce subscription plan limits and provide operational tools for support and account administration.",
])

add_heading(doc, "4. Non-goals for the MVP", 1)
add_bullets(doc, [
    "Vacation rental booking, calendars, guest communication, deposits, housekeeping, or property management.",
    "Tenant screening, lease execution, rent collection, maintenance, or landlord accounting.",
    "Automated publication to Realtor.com, Move, RAJ/MLS, or other external portals before agreements and technical requirements are confirmed.",
    "Native mobile applications; the MVP will be a responsive web application.",
    "Multiple simultaneous brokerage memberships for one agent.",
    "Custom brokerage domains, advanced website themes, paid promotion, and advanced CRM automation.",
    "Commission calculation, transaction management, conveyancing, or dispute resolution.",
    "Belize or other country localization in the Jamaica MVP.",
])

add_heading(doc, "5. User groups", 1)
add_table(doc, ["User group", "MVP purpose"], [
    ("Visitor", "Search and view listings and contact an agent without registration."),
    ("Registered consumer", "Use free account features such as favourites, saved searches, and inquiry history."),
    ("Agent", "Create and manage listing work, submit changes, share approved listings, operate a personal website, and manage inquiries."),
    ("Broker staff", "Perform delegated brokerage administration, approvals, agent management, reports, and audit review; may also be an agent."),
    ("Broker", "Control the brokerage, staff, agents, subscriptions, approvals, website, listings, and reporting; includes all staff capabilities."),
    ("SteadFast operations", "Support customers, manage billing records, monitor system and feed issues, record flags, and communicate with brokers."),
    ("SteadFast administrator", "Configure the platform, plans, permissions, internal staff, integrations, security records, and system controls."),
], [2200, 7160], font_size=9.4)
add_callout(doc, "Detailed permissions", "This PRD defines capability boundaries. The next document will be the authoritative role and permission matrix, including every create, view, edit, approve, deactivate, export, and administrative action.")

add_heading(doc, "6. Confirmed business rules", 1)
add_numbered(doc, [
    "A professional agent account requires a brokerage relationship; independent agents cannot operate in the MVP.",
    "An agent belongs to one brokerage at a time.",
    "A brokerage owns its listings, while every public listing must have an active assigned agent representative.",
    "A broker may also be an agent, and broker staff may also be agents. One person uses one account and pays once even when holding several roles.",
    "Broker staff may approve their own submissions when their brokerage grants approval permission.",
    "Listing creation, material edits, price changes, agent reassignment, removal, sold status, and rented status require brokerage approval.",
    "SteadFast operations does not approve brokerage listings or decide listing, commission, ownership, or agent disputes.",
    "Agent-to-agent sharing is advertising permission and does not require a separate broker approval after the listing is approved.",
    "A displaying agent cannot edit the shared listing and may remove it only from that agent's own website.",
    "Shared property pages show both the displaying agent and listing-owner agent; the consumer chooses whom to contact.",
    "When an agent leaves, the account and personal website remain active, but former-brokerage listings are removed from that website and public display until reassigned.",
    "All important account, listing, approval, sharing, publication, and administrative actions are retained in an audit history.",
])

add_heading(doc, "7. Functional requirements", 1)

add_heading(doc, "7.1 Accounts and authentication", 2)
add_req("FR-001", "Consumer account", "Requirement: A visitor shall be able to create, verify, sign in to, sign out of, and recover a free consumer account.", "Acceptance: A verified consumer can sign in, recover access, and reach favourites, saved searches, and inquiry history; authentication errors do not disclose whether another person's account exists.")
add_req("FR-002", "Professional identity", "Requirement: A professional shall use one identity and profile even when assigned multiple brokerage roles.", "Acceptance: Adding broker-staff or broker permissions does not create a second account or duplicate subscription for the same person.")
add_req("FR-003", "Agent application", "Requirement: An agent application shall identify the referring brokerage and remain inactive until the brokerage approves the relationship and SteadFast completes account, subscription, and required registration checks.", "Acceptance: An applicant cannot access professional listing data before activation; approval, rejection, return, and activation are timestamped and attributed.")
add_req("FR-004", "Secure session", "Requirement: The system shall provide secure session handling, password protection, account lock and recovery controls, and stronger authentication for privileged roles.", "Acceptance: Protected pages reject unauthenticated access, role changes invalidate or refresh authorization promptly, and privileged accounts can be required to use multi-factor authentication.")

add_heading(doc, "7.2 Brokerage, staff, and agent administration", 2)
add_req("FR-005", "Brokerage profile", "Requirement: A broker shall manage the company name, description, logo, contact details, office locations, branding, service areas, and public website settings.", "Acceptance: Approved profile changes appear on the brokerage website and are visible only to authorized brokerage users before publication.")
add_req("FR-006", "Staff administration", "Requirement: A broker shall invite, activate, deactivate, and assign permissions to non-agent staff and dual-role agent staff within plan limits.", "Acceptance: Deactivated staff immediately lose professional access; historical actions retain the original person's identity and role at the time of action.")
add_req("FR-007", "Agent administration", "Requirement: A broker or authorized staff member shall approve or deny agent applications, activate or deactivate brokerage access, and view the brokerage's agent roster.", "Acceptance: The roster shows status, role, subscription condition, website state, assigned listing counts, and relevant application history.")
add_req("FR-008", "Plan limits", "Requirement: The system shall enforce Broker Core capacity of 20 affiliated agents and 5 non-agent staff, Broker Growth capacity of 75 affiliated agents and 15 non-agent staff, and configurable Enterprise limits.", "Acceptance: Invitations or activations beyond a plan limit are blocked with a clear upgrade message; an agent with staff permissions counts as an agent rather than an additional non-agent staff seat.")

add_heading(doc, "7.3 Property and listing management", 2)
add_req("FR-009", "Property record", "Requirement: An authorized agent shall create a property record with structured address, geographic coordinates, property type, features, media, and reusable physical facts.", "Acceptance: Required fields are validated, coordinates can be reviewed on a map, and a property may support listing history without duplicating its stable identity.")
add_req("FR-010", "Listing draft", "Requirement: An agent shall create and save a private listing draft for an assigned brokerage-owned property.", "Acceptance: Drafts are visible only to the assigned agent and authorized brokerage users, support incomplete saves, and never appear in public or professional search.")
add_req("FR-011", "Listing content", "Requirement: A listing shall support purpose, price, currency, availability, descriptions, features, assigned representative, media ordering, visibility intent, and required declarations.", "Acceptance: The system blocks submission until mandatory fields, assigned active agent, primary image, and required brokerage declarations are complete.")
add_req("FR-012", "Media management", "Requirement: Authorized users shall upload, order, caption, replace, and remove listing images and permitted documents.", "Acceptance: Uploads enforce file type, size, access, and malware-safety controls; removed private media cannot be retrieved through a public URL.")

add_heading(doc, "7.4 Approval and listing lifecycle", 2)
add_req("FR-013", "Submit for approval", "Requirement: An agent shall submit a complete draft or material change to an approval queue with an optional message.", "Acceptance: Submission creates an immutable review version, changes the working status, notifies eligible reviewers, and prevents the pending version from replacing the approved public version.")
add_req("FR-014", "Review comparison", "Requirement: An authorized reviewer shall see the approved values, proposed values, changed fields, media changes, submitter, and submission time.", "Acceptance: The review clearly distinguishes additions, removals, and replacements and preserves the comparison after the decision.")
add_req("FR-015", "Approval decision", "Requirement: An authorized reviewer shall approve, return with comments, or reject a submission.", "Acceptance: Approval publishes the accepted version when eligible; return or rejection records a required explanation and notifies the submitter without changing the approved public version.")
add_req("FR-016", "Material changes", "Requirement: Price, public description, material property facts, media, assigned representative, visibility, withdrawal, removal, sold status, and rented status shall require approval.", "Acceptance: A material change cannot bypass the approval queue through the user interface, API, import, or administrative shortcut except a separately audited emergency control reserved for SteadFast administrators.")
add_req("FR-017", "Lifecycle states", "Requirement: The system shall support draft, pending approval, returned, rejected, approved inactive, active, changes pending, withdrawn, sold, rented, expired, and unassigned states.", "Acceptance: Only eligible active records appear publicly, and every state transition follows an allowed rule with actor, time, reason, and source recorded.")

add_heading(doc, "7.5 Public, professional, and private visibility", 2)
add_req("FR-018", "Visibility levels", "Requirement: Listings shall support private, professional-network, and public visibility subject to lifecycle eligibility.", "Acceptance: Private records are limited to authorized brokerage users; professional records are limited to authenticated eligible professionals; public records are accessible without sign-in.")
add_req("FR-019", "Publication guard", "Requirement: Public display shall require an approved active listing, active brokerage, active subscription entitlement, and active assigned agent.", "Acceptance: If any required condition becomes false, the listing is removed from public search, websites, maps, and shared displays within the defined publication-update interval.")

add_heading(doc, "7.6 Agent-to-agent sharing", 2)
add_req("FR-020", "Grant display permission", "Requirement: The assigned owner agent shall grant an eligible active agent permission to display an approved listing.", "Acceptance: The share identifies the listing, owner agent, displaying agent, grant time, status, and source and does not grant edit or approval rights.")
add_req("FR-021", "Shared display", "Requirement: An active share shall allow the listing to appear on the displaying agent's website with both agent contacts.", "Acceptance: The visitor can choose either agent; the selected agent is recorded as primary while both receive the inquiry notification.")
add_req("FR-022", "Share removal and updates", "Requirement: The owner agent may revoke a share and the displaying agent may remove it from that agent's own website; approved listing changes shall propagate automatically.", "Acceptance: Revocation, self-removal, listing changes, and status changes notify affected agents and update or remove all displays without permitting the displaying agent to alter listing content.")

add_heading(doc, "7.7 Public marketplace and maps", 2)
add_req("FR-023", "Public search", "Requirement: Visitors shall search eligible public listings using location, purpose, property type, price, bedrooms, bathrooms, and other launch filters.", "Acceptance: Search results are consistent across list and map views, support pagination or progressive loading, and exclude ineligible records.")
add_req("FR-024", "Map behaviour", "Requirement: The map shall show individual markers at close zoom and clustered results or area counts at wider zoom levels.", "Acceptance: Moving or zooming refreshes results for the visible map area, clusters expand predictably, and selecting a marker or result opens the corresponding listing summary.")
add_req("FR-025", "Property page", "Requirement: A public property page shall present approved media, facts, price, description, map, status, brokerage attribution, and relevant agent contacts.", "Acceptance: The page uses only the current approved public version and never exposes private documents, internal comments, audit data, or unpublished precise fields.")

add_heading(doc, "7.8 Agent and brokerage websites", 2)
add_req("FR-026", "Agent subdomain", "Requirement: Every active paid professional agent shall receive a persistent SteadFast subdomain with profile, branding, service areas, contacts, owned listings, and authorized shared listings.", "Acceptance: The website remains available when the agent changes brokerages, but former-brokerage listings are removed and the site accurately reflects the current professional relationship.")
add_req("FR-027", "Brokerage subdomain", "Requirement: Every active paid brokerage shall receive a branded subdomain with company profile, locations, agent directory, search, map, and eligible brokerage-owned listings.", "Acceptance: The site displays only listings owned by the brokerage and represented by active agents; property inquiries route to agents rather than a general brokerage contact action.")

add_heading(doc, "7.9 Inquiries and notifications", 2)
add_req("FR-028", "Property inquiry", "Requirement: A visitor shall submit a property or viewing inquiry without mandatory registration and select the available contact agent where more than one is shown.", "Acceptance: Required contact consent and anti-abuse checks are enforced; the inquiry stores listing version, source website, selected agent, secondary agent, time, and delivery status.")
add_req("FR-029", "Professional inbox", "Requirement: Agents shall view assigned inquiries and their status in a secure workspace.", "Acceptance: An agent can open, acknowledge, and mark an inquiry handled; unauthorized users cannot view another agent's private contact information or inquiries.")
add_req("FR-030", "Launch notifications", "Requirement: The MVP shall provide in-app and email notifications for applications, approvals, returns, rejections, listing changes, shares, revocations, inquiries, account changes, and plan-limit events.", "Acceptance: Each notification has an event source, recipient, timestamp, read state, delivery state, and link to an authorized destination; duplicate delivery is controlled.")

add_heading(doc, "7.10 Consumer account features", 2)
add_req("FR-031", "Favourites", "Requirement: A registered consumer shall add and remove eligible public listings from favourites.", "Acceptance: Favourites synchronize across signed-in sessions and retain a safe reference when a listing becomes unavailable without exposing private content.")
add_req("FR-032", "Saved searches", "Requirement: A registered consumer shall save, rename, run, and delete a search definition.", "Acceptance: Saved filters reproduce the intended search; automated alerts may remain disabled until notification policy is approved.")
add_req("FR-033", "Inquiry history", "Requirement: A registered consumer shall view inquiries submitted while signed in.", "Acceptance: History shows property, selected agent, submission date, and status information approved for consumers without exposing internal professional notes.")

add_heading(doc, "7.11 Subscription and billing administration", 2)
add_req("FR-034", "Professional plan", "Requirement: The system shall support a US$50 monthly Professional Agent plan and record whether the seat is self-paid or brokerage-funded.", "Acceptance: Entitlements depend on active plan status, and a dual-role person is represented by one professional seat.")
add_req("FR-035", "Broker plans", "Requirement: The system shall support Broker Core at US$150 monthly, Broker Growth at US$300 monthly, and configurable Enterprise pricing and limits.", "Acceptance: Plan assignment controls capacity and brokerage features without silently deleting users or listings when a limit is exceeded.")
add_req("FR-036", "Pilot billing records", "Requirement: SteadFast operations shall create invoice records, record payment and adjustment status, view subscription history, and manage grace or suspension status.", "Acceptance: Every billing change is audited; the MVP may use administrator-recorded payments until an approved payment gateway is integrated before public launch.")

add_heading(doc, "7.12 SteadFast operations and administration", 2)
add_req("FR-037", "Customer support", "Requirement: Operations staff shall search authorized customer accounts, view non-sensitive support context, record issues, manage billing assistance, and communicate account actions.", "Acceptance: Support actions are permission-limited and audited, and operations staff cannot approve brokerage listings or resolve professional disputes.")
add_req("FR-038", "Flags", "Requirement: Operations staff shall record and monitor a flagged listing or account and notify the responsible brokerage.", "Acceptance: A flag records source, category, evidence reference, brokerage notification, response, status, and timeline without automatically suspending listing content.")
add_req("FR-039", "Platform administration", "Requirement: Authorized SteadFast administrators shall manage internal staff permissions, plans, limits, controlled configuration, integration credentials, security events, and emergency actions.", "Acceptance: High-risk changes require explicit permission, confirmation, reason, and immutable audit records; secrets are never displayed after initial secure entry.")

add_heading(doc, "7.13 Audit and history", 2)
add_req("FR-040", "Audit record", "Requirement: Security-sensitive and business-significant actions shall create an append-only audit event.", "Acceptance: Each event records actor, effective role, brokerage context, action, target, time, source, request reference, and safe before/after summary; ordinary users cannot edit or delete events.")
add_req("FR-041", "Visible history", "Requirement: Agents and authorized brokerage users shall view the history relevant to their listings, submissions, decisions, shares, and assignments.", "Acceptance: History respects brokerage and role boundaries and excludes internal security data, secrets, and unrelated customer information.")

add_heading(doc, "7.14 API and integration readiness", 2)
add_req("FR-042", "Canonical records", "Requirement: Core property, listing, professional, brokerage, media, status, and distribution data shall use stable internal identifiers and versioned records.", "Acceptance: External identifiers never replace internal identity, and a listing change can be traced across approval and future distribution attempts.")
add_req("FR-043", "Integration boundary", "Requirement: External portal and MLS adapters shall remain separate from core listing ownership and approval logic.", "Acceptance: Disabling or changing one adapter does not prevent SteadFast from managing and publishing its own approved listings.")
add_req("FR-044", "Distribution eligibility", "Requirement: Future distribution shall evaluate brokerage authorization, listing permission, plan entitlement, destination rules, and approved status before export.", "Acceptance: No external send can occur without a recorded eligible result; delivery attempts, responses, errors, updates, and removals are retained.")

add_heading(doc, "8. Non-functional requirements", 1)
add_table(doc, ["Area", "MVP requirement"], [
    ("Usability", "Responsive, plain-language interface; core agent listing submission should require no formal training."),
    ("Accessibility", "Keyboard-accessible interactions, visible focus, labelled forms, sufficient contrast, meaningful errors, and practical WCAG 2.2 AA alignment."),
    ("Performance", "Public pages should render useful content quickly on typical Jamaican mobile connections; search and map interactions must remain responsive at pilot scale."),
    ("Availability", "Separate preview and production environments, health monitoring, actionable error reporting, and documented incident ownership."),
    ("Security", "Server-side authorization, brokerage data isolation, least privilege, protected secrets, secure uploads, rate limits, audit records, and privileged multi-factor authentication."),
    ("Privacy", "Data minimization, clear notices and consent, controlled retention, export/correction/deletion workflows where applicable, and production logging that excludes sensitive values."),
    ("Reliability", "Versioned migrations, automated tests, backups, recovery verification, idempotent event processing, and controlled deployment rollback."),
    ("Portability", "Private Git repository, documented configuration, standard PostgreSQL data, migrations, storage abstraction, and deployable Node or container configuration."),
    ("Observability", "Request correlation, structured logs, system metrics, error monitoring, audit events, and integration-delivery status."),
], [1900, 7460], font_size=9.1)

add_heading(doc, "9. Analytics and success measures", 1)
add_bullets(doc, [
    "Application-to-activation conversion by brokerage and professional role.",
    "Time from activation to first draft, first submission, first approval, and first public listing.",
    "Approval turnaround, return rate, rejection rate, and most common validation failures.",
    "Weekly and monthly active professionals, active brokerages, public listings, shared displays, and agent website usage.",
    "Search activity, map usage, property views, inquiries, selected contact agent, and inquiry acknowledgement.",
    "Paid conversion, recurring revenue, failed payments, grace status, retention, and churn.",
    "Support volume, repeated issue categories, resolution time, system incidents, and customer satisfaction.",
])

add_heading(doc, "10. Release acceptance criteria", 1)
add_numbered(doc, [
    "All launch roles can authenticate and access only authorized workspaces and brokerage data.",
    "An agent can complete the full draft-to-approved-public listing journey using test data.",
    "Every material change and terminal status follows the approval workflow and remains historically traceable.",
    "No unapproved, inactive, unassigned, or unauthorized listing appears in public search, maps, subdomains, or shared displays.",
    "Agent departure removes former-brokerage listings from the agent website and public display until reassigned without deleting required history.",
    "Agent sharing preserves ownership, prevents recipient editing, displays both contacts, and produces the required notifications.",
    "Visitors can search, use the map, view a property, choose an agent, and submit an inquiry without registration.",
    "Agent and brokerage websites reflect approved changes and removals within the agreed publication-update interval.",
    "Plan limits, billing status, and grace rules behave predictably without corrupting listing or account data.",
    "Automated tests cover critical permission, approval, publication, inquiry, billing-entitlement, and audit scenarios.",
    "Backups, restore procedure, monitoring, incident contacts, and rollback process are documented and tested before the external pilot.",
    "A security threat model is completed before external pilot access and all high-risk findings are resolved or formally accepted.",
])

add_heading(doc, "11. Delivery sequence", 1)
add_table(doc, ["Build stage", "Included outcome"], [
    ("1. Foundation", "Repository, environments, authentication, brokerage model, roles, plan records, audit foundation, and documentation structure."),
    ("2. Listing operations", "Property, media, drafts, validation, approval versions, comparisons, decisions, status rules, and reassignment."),
    ("3. Public discovery", "Publication rules, search, filters, property pages, maps, agents, brokerages, and inquiries."),
    ("4. Websites and sharing", "Agent and brokerage subdomains, display permissions, dual-agent contacts, notifications, and automatic updates."),
    ("5. Commercial operations", "Plan limits, invoice records, payment status, support workspace, flags, administration, and pilot analytics."),
    ("6. Launch hardening", "Security review, performance, accessibility, backups, recovery, monitoring, operational runbooks, and pilot readiness."),
], [2200, 7160], font_size=9.2)

add_heading(doc, "12. Dependencies", 1)
add_bullets(doc, [
    "Confirmed SteadFast legal entity, product owner, and authority to approve scope and pricing.",
    "Founding brokerages and agents available for workflow validation and pilot testing.",
    "Private company Git repository and controlled access for development staff.",
    "Vercel, Supabase, domain or subdomain, email, map, monitoring, and file-scanning accounts when implementation reaches those services.",
    "Legal review of privacy, subscription, acceptable-use, content-rights, broker responsibility, pilot, and data-processing terms.",
    "Payment provider selection before automated production billing.",
    "Written data and distribution agreements before any external MLS or portal publication.",
])

add_heading(doc, "13. Open decisions not blocking foundation work", 1)
add_table(doc, ["Decision", "Working MVP position"], [
    ("Payment gateway", "Use administrator-recorded pilot payment status; select and integrate the approved gateway before public launch."),
    ("JMD billing", "Maintain USD product prices as the current reference and define a stable JMD price-review policy before invoicing locally."),
    ("SMS or WhatsApp", "Use in-app and email notifications in MVP; add other channels after consent, provider, cost, and delivery policy are approved."),
    ("Listing expiry", "Support an expiration state; exact default duration and renewal workflow require broker pilot feedback."),
    ("Approximate location", "Use verified address coordinates in MVP; define broker-controlled approximate public display only if privacy demand is confirmed."),
    ("Consumer alerts", "Save searches in MVP; activate automated alerts after frequency and consent rules are approved."),
    ("International feeds", "Build canonical data and adapter boundaries now; activate destinations only after written authorization."),
], [2600, 6760], font_size=9.2)

add_heading(doc, "14. Glossary", 1)
add_table(doc, ["Term", "Meaning"], [
    ("Brokerage", "The professional company or dealer organization that owns listings and manages agents and staff in SteadFast."),
    ("Broker", "The primary brokerage account holder with all broker-staff capabilities and company-level control."),
    ("Broker staff", "A delegated brokerage user who may manage agents, approvals, reporting, or administration; may also be an agent."),
    ("Owner agent", "The assigned representative responsible for a brokerage-owned listing and able to grant display permission."),
    ("Displaying agent", "An agent authorized to advertise another agent's approved listing without ownership or editing rights."),
    ("Property", "The physical real-world asset and stable facts that may support one or more listing periods."),
    ("Listing", "A brokerage-controlled offer to sell or rent a property with price, representation, publication, and lifecycle status."),
    ("Approved version", "The immutable listing version accepted by an authorized brokerage reviewer."),
    ("Canonical record", "SteadFast's authoritative internal data, independent of external portal identifiers or formats."),
], [2300, 7060], font_size=9.3)

add_heading(doc, "15. Document governance", 1)
add_body(doc, "This PRD is the product scope baseline. Changes to confirmed business rules, MVP inclusions, or acceptance criteria require a recorded decision and version update. More detailed technical documents may refine implementation but must not silently change product behaviour. The next document is the Roles and Permissions Matrix.")

settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "SteadFast MVP Product Requirements"
doc.core_properties.subject = "Jamaica pilot product requirements and acceptance baseline"
doc.core_properties.author = "SteadFast"
doc.core_properties.keywords = "SteadFast, MVP, PRD, Jamaica, real estate, brokers, agents, listings"
doc.save(str(OUTPUT))
print(OUTPUT)
