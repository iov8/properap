from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_steadfast_roles_permissions as shared_docx

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SteadFast_Listing_Workflow_Specification_v0.1.md"
OUTPUT = ROOT / "docs" / "SteadFast_Listing_Workflow_Specification_v0.1.docx"
TEMPLATE = ROOT / "docs" / "SteadFast_MVP_Product_Requirements_v0.1.docx"

from build_steadfast_roles_permissions import (
    add_page_number,
    set_cell_shading,
    set_font,
    set_repeat_header,
    set_table_geometry,
)

shared_docx.TABLE_INDENT = 120

BLUE = "1F4E78"
DARK_BLUE = "17365D"
DARK = "1F2937"
MID_GRAY = "667085"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "EEF5FA"


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_inline(paragraph, text, size=11, color=DARK):
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`") and token.endswith("`"):
            set_font(paragraph.add_run(token[1:-1]), size=max(9, size - 0.5), color=DARK_BLUE, name="Consolas")
        else:
            set_font(paragraph.add_run(token), size=size, color=color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def table_widths(count):
    choices = {
        2: [2200, 7160],
        3: [1800, 3000, 4560],
        4: [1500, 2200, 1700, 3960],
        5: [1200, 1700, 1400, 1600, 3460],
    }
    if count in choices:
        return choices[count]
    base = 9360 // count
    return [base] * (count - 1) + [9360 - base * (count - 1)]


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, header, size=9.2, color=DARK_BLUE)
        for run in p.runs:
            run.bold = True
    for row_index, values in enumerate(rows[1:], start=1):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            if row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, value, size=8.8)
    set_table_geometry(table, table_widths(len(headers)))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        parts = [item.strip() for item in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", item) for item in parts):
            rows.append(parts)
        index += 1
    return rows, index


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("STEADFAST  |  LISTING WORKFLOW SPECIFICATION"), size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Workflow Baseline  |  Version 0.1  |  Page "), size=9, color=MID_GRAY)
    add_page_number(footer)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("PRODUCT AND ENGINEERING SPECIFICATION"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("SteadFast Listing Workflow"), size=28, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("Approval-safe listing operations for the Jamaica MVP"), size=14, color=MID_GRAY)
    metadata = (
        ("Document owner", "SteadFast Product Management"),
        ("Status", "Planning baseline for product, engineering, design, testing, and operations"),
        ("Version", "0.1"),
        ("Date", "16 July 2026"),
        ("Depends on", "MVP Product Requirements and Roles and Permissions Matrix v0.1"),
    )
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=DARK_BLUE)
        set_font(p.add_run(value), size=10.5, color=DARK)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("Core control: "), size=11, bold=True, color=BLUE)
    set_font(p.add_run("A pending change never replaces the approved public version. Publication changes only after an authorized, atomic approval and a fresh eligibility check."), size=11, color=DARK)
    doc.add_paragraph()


def build():
    doc = Document(TEMPLATE)
    clear_body(doc)
    configure_styles(doc)
    add_masthead(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("**Version:") or line.startswith("**Prepared:") or line.startswith("**Status:") or line.startswith("**Applies to:"):
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            add_inline(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = "SteadFast Listing Workflow Specification v0.1"
    props.subject = "MVP listing creation, approval, publication, sharing, reassignment, and lifecycle workflow"
    props.author = "SteadFast"
    props.keywords = "SteadFast, listing workflow, approvals, property, brokerage, agent"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
