from pathlib import Path
import runpy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)
PLAN_PATH = DOCS / "SteadFast_Product_and_Technical_Plan_v0.1.docx"
OUTPUT = DOCS / "SteadFast_Business_Proposal_v0.1.docx"

# Reuse the established SteadFast Word design system and helpers.
ns = runpy.run_path(str(ROOT / "tools" / "build_steadfast_plan.py"))
doc = Document(str(PLAN_PATH))

# Clear the planning-document body while retaining its styles and section setup.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

BLUE = ns["BLUE"]
DARK_BLUE = ns["DARK_BLUE"]
DARK = ns["DARK"]
MID_GRAY = ns["MID_GRAY"]
PALE_BLUE = ns["PALE_BLUE"]
LIGHT_BLUE = ns["LIGHT_BLUE"]

set_font = ns["set_font"]
add_heading = ns["add_heading"]
add_body = ns["add_body"]
add_bullets = ns["add_bullets"]
add_numbered = ns["add_numbered"]
add_callout = ns["add_callout"]
add_table = ns["add_table"]
add_page_number = ns["add_page_number"]

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Narrative-proposal preset.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Proposal running header/footer.
header = section.header
hp = header.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("STEADFAST  |  BUSINESS PROPOSAL")
set_font(hr, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Confidential Planning Draft  |  Version 0.1  |  Page ")
set_font(fr, size=9, color=MID_GRAY)
add_page_number(fp)


def centered(text, size=11, bold=False, color=DARK, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_lead(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_font(r, size=12, bold=True, color=DARK_BLUE)
    return p


# Proposal centerpiece cover.
centered("STEADFAST", size=12, bold=True, color=BLUE, after=22)
centered("A Modern Real Estate Platform", size=26, bold=True, color=DARK_BLUE, after=5)
centered("for Jamaica and the Caribbean", size=19, bold=True, color=BLUE, after=12)
centered(
    "Business proposal for a broker-controlled listing, marketing, website, collaboration, and international-distribution platform",
    size=12,
    color=MID_GRAY,
    after=24,
)

add_table(doc, ["Proposal", "Details"], [
    ("Prepared by", "SteadFast"),
    ("Initial market", "Jamaica"),
    ("Expansion direction", "Belize and selected Caribbean and Central American markets"),
    ("Document status", "Client discussion draft - Version 0.1"),
    ("Date", "16 July 2026"),
], [2400, 6960], header_fill=LIGHT_BLUE, font_size=10)

add_callout(
    doc,
    "Core proposition",
    "SteadFast gives brokerages and agents one easy system to create, approve, market, share, and distribute property listings while preserving brokerage control and professional accountability.",
)

add_heading(doc, "1. Executive summary", 1)
add_lead("SteadFast is designed to become the daily operating platform for Jamaican real-estate brokerages and agents.")
add_body(doc, "The platform combines listing management, broker approval, agent collaboration, branded public websites, map-based property discovery, customer inquiries, subscriptions, audit history, and integration-ready property data in one cloud service. It is intentionally designed for novice users so that an agent can create and manage a listing without formal software training.")
add_body(doc, "Jamaica is the launch market. Once the operating model, product, and commercial relationships are proven, the same platform can be localized for selected Caribbean and Central American markets. Belize is a logical early expansion candidate, followed by other markets chosen through local brokerage partnerships and market validation.")
add_body(doc, "The recommended launch model keeps browsing free, charges real-estate professionals for productive tools, and charges brokerages for company-level control, branding, administration, and reporting. International distribution through Realtor.com and Move-operated channels should be presented as a planned, agreement-dependent capability rather than a guaranteed day-one feature.")

add_heading(doc, "2. The market problem", 1)
add_body(doc, "Real-estate professionals need more than a public property portal. They need a reliable operating system that reflects how brokerages actually work: agents create listings, brokers control quality and accountability, multiple agents may market the same property, and customers expect modern search, maps, mobile access, and quick communication.")
add_bullets(doc, [
    "Existing tools may be difficult to learn, slow to use, or poorly aligned with local brokerage workflows.",
    "Listings, approvals, website updates, agent relationships, and customer inquiries are often spread across separate tools and informal messaging.",
    "Brokerages need ownership and oversight without preventing agents from marketing effectively.",
    "Agents need a professional web presence without separately purchasing and maintaining a website.",
    "International buyers need accurate, visually strong listings and a clear path to a qualified local agent.",
    "Future portal and MLS connections require consistent data, permissions, status history, and distribution controls from the beginning.",
])

add_heading(doc, "3. The SteadFast solution", 1)
add_body(doc, "SteadFast will provide a secure cloud platform centered on the brokerage. Each brokerage manages its company, staff, agents, approvals, subscriptions, websites, and listing portfolio. Each professional receives a role-appropriate workspace, and every public visitor receives a modern property-search experience.")
add_callout(doc, "Product position", "SteadFast is not simply another advertising website. It is a listing operating system, professional collaboration network, branded website service, and controlled distribution platform.")

add_heading(doc, "4. Who the platform serves", 1)
add_table(doc, ["User", "What SteadFast provides"], [
    ("Visitors", "Free search, maps, property pages, agent discovery, and inquiry forms without mandatory registration."),
    ("Registered consumers", "Free favourites, saved searches, inquiry history, and future property alerts."),
    ("Agents", "Listing creation, approval submissions, personal website, sharing, inquiries, notifications, profile, and subscription tools."),
    ("Broker staff", "Approval queues, agent administration, reassignment, portfolio oversight, flags, reports, and audit history; staff may also be agents."),
    ("Brokers", "All staff capabilities plus company control, staff permissions, agent applications, billing, integrations, reporting, and brokerage website."),
    ("SteadFast operations", "Customer support, account and billing assistance, technical issue handling, feed monitoring, and broker notifications."),
    ("SteadFast administrators", "Platform configuration, plans, internal permissions, integration management, security logs, and system oversight."),
], [2000, 7360], font_size=9.4)

add_heading(doc, "5. Complete platform capabilities", 1)
add_heading(doc, "5.1 Listing creation and lifecycle", 2)
add_bullets(doc, [
    "Create residential sale, long-term rental, commercial, land, and development listings.",
    "Capture structured property details, address, map coordinates, descriptions, pricing, features, photographs, documents, and agent representation.",
    "Save private drafts and submit complete listings for brokerage approval.",
    "Maintain statuses such as draft, pending approval, active, changes pending, returned, withdrawn, sold, rented, expired, and unassigned.",
    "Preserve every submission, revision, approval, return, status change, and responsible user in a permanent audit history.",
])

add_heading(doc, "5.2 Broker-controlled approval", 2)
add_bullets(doc, [
    "Broker or authorized staff approval is required before a listing becomes public.",
    "Price changes, major edits, representative changes, removal, sold status, and rented status require approval.",
    "Reviewers see the exact difference between the approved version and the proposed change.",
    "Reviewers may approve, return with comments, or reject a submission.",
    "Authorized broker staff may approve their own submissions when the brokerage permits it.",
])

add_heading(doc, "5.3 Brokerage ownership and agent continuity", 2)
add_bullets(doc, [
    "The brokerage owns its listings while an active agent serves as the required public representative.",
    "An agent may belong to one brokerage at a time during the initial release.",
    "When an agent leaves, the account and personal website remain active, but former-brokerage listings are removed from that agent's public website and taken out of public display until reassigned.",
    "The same account may later join a new brokerage without losing identity or subscription history.",
    "A person who is both broker staff and an agent uses one account and pays once.",
])

add_heading(doc, "5.4 Agent-to-agent advertising", 2)
add_bullets(doc, [
    "An agent may allow another agent to display an approved listing without requesting a new broker approval.",
    "Sharing grants display permission only; it never transfers ownership or editing authority.",
    "The shared property page shows both the displaying agent and the listing-owner agent, and the consumer chooses whom to contact.",
    "Both professionals receive inquiry notifications, with the selected agent recorded as the primary contact.",
    "Removal, revocation, and approved listing changes trigger notifications and remain in the audit history.",
])

add_heading(doc, "5.5 Agent and brokerage websites", 2)
add_bullets(doc, [
    "Every active agent receives a persistent branded subdomain website with profile, biography, service areas, contact methods, listings, search, and map tools.",
    "Every brokerage receives a branded website showing all approved public listings it owns, its active agents, and its company profile.",
    "Shared listings may appear on an agent website when display permission is active.",
    "Property inquiries are directed to agents. Brokerage pages build trust and attribution but do not replace the need for agent service.",
    "Future paid plans may support custom domains, advanced branding, featured listings, and lead-routing rules.",
])

add_heading(doc, "5.6 Public discovery and customer experience", 2)
add_bullets(doc, [
    "Mobile-friendly public homepage, property search, filters, list view, and synchronized map view.",
    "Close zoom shows individual markers; wider views group results into numbered clusters and geographic areas.",
    "Detailed property pages present media, features, location, status, and relevant agent choices.",
    "Visitors may contact an agent or request a viewing without creating an account.",
    "Registered users may save favourites, searches, and inquiry history, with alerts introduced later.",
])

add_heading(doc, "5.7 Professional operations", 2)
add_bullets(doc, [
    "Role-specific dashboards for agents, broker staff, brokers, SteadFast operations, and administrators.",
    "Agent applications require broker referral, broker approval, and platform account completion.",
    "Brokers can appoint staff, manage permissions, deactivate professional access, and reassign listings.",
    "Invoices, subscription status, payment history, plan eligibility, and service support are managed centrally.",
    "SteadFast may monitor flagged listings and notify the responsible brokerage, while listing disputes and content decisions remain with the broker.",
])

add_heading(doc, "5.8 Integration-ready data platform", 2)
add_body(doc, "SteadFast will maintain its own authoritative listing record and a separate integration layer. The data model will use RESO-aligned concepts where practical so property, member, office, and media data can be mapped consistently to authorized external systems.")
add_bullets(doc, [
    "Documented application programming interfaces for approved partners and future applications.",
    "External identifiers, eligibility rules, validation results, delivery status, errors, and update history for every distribution channel.",
    "Channel-specific adapters so one portal can change without disrupting the core listing workflow.",
    "Permission and subscription controls that determine which listings may be distributed and where.",
    "A foundation for future MLS, association, portal, analytics, CRM, and marketing integrations.",
])

add_heading(doc, "6. International buyer strategy", 1)
add_body(doc, "SteadFast will help Jamaican professionals present local inventory to overseas buyers through high-quality public pages, accurate map data, structured property records, and qualified agent contacts. The preferred distribution direction is Realtor.com international exposure and other Move-operated or approved partner channels.")
add_callout(doc, "Commercial dependency", "Move operates Realtor.com, but a compatible data feed does not itself create the right to publish there. Distribution requires written authorization, feed acceptance, contractual terms, and ongoing compliance with the destination's rules.")
add_numbered(doc, [
    "Build SteadFast's canonical listing model and export controls using RESO-aligned concepts.",
    "Confirm the correct business and technical contact for Realtor.com International or the applicable Move distribution program.",
    "Obtain written data-feed, branding, content-rights, update-frequency, and compliance requirements.",
    "Pilot a controlled set of approved Jamaican listings and monitor acceptance, errors, updates, and removals.",
    "Offer international distribution only in the plans and markets for which SteadFast has confirmed rights and reliable operations.",
])

add_heading(doc, "7. Expansion strategy", 1)
add_table(doc, ["Stage", "Market direction", "Purpose"], [
    ("1. Launch", "Jamaica", "Prove broker adoption, approval workflows, public discovery, websites, billing, and listing quality."),
    ("2. Replicate", "Belize", "Test English-language regional expansion with country-specific fields, currency, locations, and professional rules."),
    ("3. Regional portfolio", "Selected Caribbean and Central American markets", "Enter markets one at a time using local brokerage partners, localization, and confirmed distribution rights."),
], [1200, 2100, 6060], font_size=9.2)
add_body(doc, "Expansion should be market-led rather than map-led. A country is ready only when SteadFast has an anchor brokerage or association relationship, a local operating and legal assessment, location and currency support, translated terminology where required, a payment method, and a viable customer-support model.")

add_heading(doc, "8. Business model", 1)
add_body(doc, "Consumers create market demand and should remain free. Revenue should come from professionals and brokerages that receive measurable productivity, control, branding, collaboration, and distribution value.")
add_bullets(doc, [
    "Free public access for visitors and registered consumers.",
    "Paid professional subscriptions for agents and broker staff.",
    "Paid brokerage plans for company administration, website, approvals, staff capacity, reporting, and integration controls.",
    "One subscription charge per person even when that person holds several professional roles.",
    "Optional future revenue from custom domains, premium website themes, advanced analytics, promoted inventory, lead-routing services, data exports, and authorized international distribution.",
])

add_heading(doc, "9. Pricing recommendation", 1)
add_lead("Use US$50 per month as the current planning price for a regular professional agent seat.")
add_body(doc, "The US$50 monthly price positions SteadFast as a professional business tool while remaining substantially below a US$200 monthly fee. The paid pilot should test whether agents consistently use the listing workspace, personal website, collaboration tools, inquiries, and international-marketing capabilities. An annual prepayment option can be introduced after the pilot confirms retention and support costs.")
add_table(doc, ["Plan", "Recommended launch price", "Included value"], [
    ("Visitor / registered consumer", "Free", "Search, maps, property pages, agent contact, favourites, saved searches, and inquiry history."),
    ("Professional seat", "US$50 monthly", "One agent or broker-staff account, listing tools when authorized, personal website, sharing, inquiries, notifications, and support."),
    ("Broker Core", "US$150 monthly or US$1,500 yearly", "Broker account, brokerage website, approvals, company controls, reporting, and up to five paid staff seats funded through the company plan; agents subscribe separately."),
    ("Broker Growth", "US$300 monthly or US$3,000 yearly", "Up to twenty company-funded staff seats, advanced reporting, priority onboarding, and integration/distribution controls when available; agents subscribe separately."),
    ("Enterprise / association", "Custom", "Large networks, negotiated support, migration, custom integration, service commitments, and volume pricing."),
], [1900, 1900, 5560], font_size=9.0)
add_callout(doc, "Billing principle", "A person is never charged twice because they are both an agent and broker staff. A seat may be paid personally or funded by the brokerage, but it represents one active professional subscription.")
add_body(doc, "For Jamaican customers, SteadFast may display and invoice an approved JMD price list while using the USD figures as internal reference. Local prices should not fluctuate every day; they should be reviewed on a predictable schedule and state whether taxes and payment-processing charges are included.")

add_heading(doc, "10. Illustrative first-year revenue", 1)
add_body(doc, "The following example is a planning scenario, not a forecast. It uses the stated early interest of approximately 200 professionals and assumes ten Broker Core subscriptions.")
add_table(doc, ["Revenue source", "Illustrative volume", "Annual value"], [
    ("Professional monthly seats", "200 x US$50 x 12", "US$120,000"),
    ("Broker Core monthly plans", "10 x US$150 x 12", "US$18,000"),
    ("Illustrative annual recurring revenue", "Before discounts, taxes, refunds, and processing", "US$138,000"),
], [3000, 3860, 2500], font_size=9.6)
add_body(doc, "This scenario assumes every interested professional becomes a paying and retained customer, which is unlikely at launch. The real model should therefore be validated through customer interviews, a paid pilot, conversion tracking, support costs, and churn data before the revenue figure is treated as a forecast.")

add_heading(doc, "11. Why SteadFast can win", 1)
add_bullets(doc, [
    "Designed around Jamaican brokerage operations instead of forcing local users into a foreign workflow.",
    "Easy enough for novice users, reducing training and support burden.",
    "Broker-controlled quality with transparent approvals and complete history.",
    "Agent websites and brokerage websites included in the same service.",
    "Collaboration that lets agents advertise each other's approved listings without losing ownership or accountability.",
    "A modern consumer experience with maps, mobile design, and direct access to professional agents.",
    "Integration-ready records and APIs designed for future portal and MLS connections.",
    "A regional architecture capable of supporting country, currency, language, location, tax, and policy differences.",
])

add_heading(doc, "12. Go-to-market plan for Jamaica", 1)
add_numbered(doc, [
    "Recruit a founding advisory group of interested brokers and experienced agents.",
    "Demonstrate clickable product workflows and confirm the minimum launch feature set.",
    "Launch a private pilot with a small number of brokerages, real inventory, and controlled public visibility.",
    "Measure time to first listing, approval turnaround, listing completeness, inquiry response, active usage, support requests, and willingness to pay.",
    "Convert the pilot to founding annual subscriptions with time-limited pricing protection.",
    "Publish verified brokerage and agent success stories, then expand through referrals and professional associations.",
    "Begin international-feed onboarding only after listing quality, rights, removal processing, and monitoring are reliable.",
])

add_heading(doc, "13. Delivery roadmap", 1)
add_table(doc, ["Phase", "Commercial outcome"], [
    ("1. Foundation", "Secure accounts, brokerage structure, roles, permissions, audit foundation, repository, environments, and documentation."),
    ("2. Listing operations", "Property records, media, drafts, approvals, comparisons, statuses, brokerage ownership, and reassignment."),
    ("3. Public marketplace", "Search, filters, property pages, maps, agent profiles, brokerage profiles, and inquiries."),
    ("4. Websites and sharing", "Agent and brokerage subdomains, advertising permissions, contact choice, and notifications."),
    ("5. Commercial operations", "Plans, subscriptions, invoices, customer support, flag monitoring, analytics, and administration."),
    ("6. External distribution", "Authorized RAJ/MLS and Realtor.com/Move adapters, validation, delivery monitoring, and tier controls."),
    ("7. Regional expansion", "Country configuration, localization, currency, payment, legal review, and partner onboarding."),
], [2200, 7160], font_size=9.3)

add_heading(doc, "14. Key dependencies and risk controls", 1)
add_bullets(doc, [
    "International publication depends on separate portal agreements and must not be sold as guaranteed until confirmed in writing.",
    "Brokerages must have the legal right to provide every listing and media item for the intended display channels.",
    "SteadFast should launch with clear subscription, refund, privacy, data-processing, acceptable-use, and content-responsibility terms.",
    "Security, multi-broker data isolation, backups, monitoring, and incident procedures must be implemented from the beginning.",
    "Pricing must be tested with actual decision-makers; interest does not equal paid adoption.",
    "Each new country requires its own professional, tax, payment, privacy, mapping, language, and data-distribution assessment.",
    "Support capacity and onboarding quality must scale with the number of professionals and listings.",
])

add_heading(doc, "15. Success measures", 1)
add_bullets(doc, [
    "A new agent can create and submit a complete listing without training.",
    "A broker can review, compare, approve, return, and trace every important change.",
    "No unapproved or unassigned listing can appear publicly.",
    "Public visitors can search Jamaican property on a map and contact an agent without registering.",
    "Agent and brokerage websites update automatically when approvals, changes, shares, and removals occur.",
    "Pilot customers actively use the platform, convert to paid plans, and recommend it to peers.",
    "External feed errors and removals are detected, recorded, and resolved within defined service targets.",
    "The platform can add a second country without rebuilding the core listing system.",
])

add_heading(doc, "16. Recommended client decision", 1)
add_lead("Approve a Jamaica-first paid pilot and treat regional expansion and international distribution as gated growth stages.")
add_body(doc, "The immediate objective is not to build every regional feature at once. It is to prove that Jamaican brokers and agents will adopt and pay for a simpler, better listing workflow; that approved inventory can produce a strong public experience; and that SteadFast can operate the service reliably. Once that foundation is demonstrated, international distribution and country expansion become credible growth engines rather than speculative promises.")
add_numbered(doc, [
    "Confirm the founding pilot brokerages and the professionals authorized to provide weekly product feedback.",
    "Approve the launch scope and the recommended founding prices for validation.",
    "Begin the local prototype, preview deployment, technical documentation, and private Git repository.",
    "Open formal discussions with RAJ/MLS and the appropriate Realtor.com/Move distribution contact.",
    "Review pilot evidence before authorizing production launch or the first expansion market.",
])

add_heading(doc, "17. Reference direction", 1)
add_bullets(doc, [
    "Move, Inc. and Realtor.com relationship: https://www.move.com/",
    "RESO Web API and authorization overview: https://www.reso.org/reso-web-api/",
    "RESO Data Dictionary and international data concepts: https://www.reso.org/knowledge-base/reso-data-dictionary-faq/",
    "Real Estate Board of Jamaica: https://reb.gov.jm/",
    "Realtors Association of Jamaica: https://realtorsjamaica.org/",
])
add_body(doc, "Pricing, revenue examples, market sequence, and delivery timing in this proposal are planning recommendations and must be validated before being represented as contractual commitments.")

# Ensure the document opens at the beginning and updates fields in Word.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "SteadFast Business Proposal"
doc.core_properties.subject = "Jamaica-first real-estate listing and marketing platform with regional expansion"
doc.core_properties.author = "SteadFast"
doc.core_properties.keywords = "SteadFast, Jamaica, real estate, SaaS, brokers, agents, listings, proposal"
doc.save(str(OUTPUT))
print(OUTPUT)
