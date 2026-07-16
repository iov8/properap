from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "SteadFast_Product_and_Technical_Plan_v0.1.docx"

BLUE = "1F4E78"
DARK_BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
WHITE = "FFFFFF"
GREEN = "287D3C"
AMBER = "8A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=None, color=DARK, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MID_GRAY)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_font(r)


def add_callout(doc, label, text, fill=PALE_BLUE, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    set_table_geometry(table, widths, indent=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    # Major sections flow continuously; heading keep-with-next rules prevent
    # orphaned section titles without creating large blank areas.
    return None


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.167

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("STEADFAST  |  PRODUCT & TECHNICAL PLAN")
set_font(hr, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Planning Draft  |  Version 0.1  |  Page ")
set_font(fr, size=9, color=MID_GRAY)
add_page_number(fp)

# First-page masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("PRODUCT & TECHNICAL PLAN")
set_font(r, size=10, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("SteadFast")
set_font(r, size=30, bold=True, color=DARK_BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Cloud listing and property marketing platform for Jamaican real-estate brokers and agents")
set_font(r, size=14, color=MID_GRAY)

metadata = [
    ("Document", "Product requirements and technical planning baseline"),
    ("Version", "0.1 - Planning Draft"),
    ("Prepared", "July 2026"),
    ("Status", "Confirmed decisions plus identified future decisions"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(value)
    set_font(r)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    doc,
    "Purpose",
    "This document records the agreed product model, workflows, application areas, engineering approach, security baseline, and delivery process for the first SteadFast implementation.",
)

add_heading(doc, "1. Executive summary", 1)
add_body(
    doc,
    "SteadFast will be a multi-tenant cloud platform for Jamaican real-estate brokers and agents to create, approve, manage, advertise, share, and distribute property listings. The brokerage is the root of the professional organization. Brokerages own approved listings, while each published listing must have an active agent representative."
)
add_body(
    doc,
    "The platform will also provide a public property-search experience, individual agent websites, brokerage websites, internal SteadFast operations tools, and a full administrative workspace. The design will remain simple enough for novice users while preserving strict permissions, complete change history, and future compatibility with industry listing feeds."
)

add_heading(doc, "2. Product goals", 1)
add_bullets(doc, [
    "Give agents a fast, understandable way to create and manage listings without formal software training.",
    "Give brokers direct control over publication, material changes, agent membership, listing assignment, and company presentation.",
    "Provide polished public property discovery through list search, map search, agent websites, and brokerage websites.",
    "Support agent-to-agent listing advertising without duplicating or transferring listing ownership.",
    "Create a standards-ready data foundation for RAJ/MLS and Realtor.com International/Move distribution when commercial authorization is available.",
    "Maintain a portable, secure, documented codebase that future SteadFast IT staff can operate and move between hosting providers.",
])

add_heading(doc, "3. Initial scope", 1)
add_table(doc, ["Included in first product", "Deferred or future"], [
    ("Residential sales and long-term rentals", "Vacation and short-term rentals"),
    ("Commercial properties, land, and developments", "Tenant management and rent collection"),
    ("Listing creation, broker approval, and change history", "Maintenance, inspections, and lease administration"),
    ("Public search, maps, inquiries, and agent selection", "Advanced transaction and commission accounting"),
    ("Agent and brokerage websites", "Native mobile applications"),
    ("Sharing for advertising", "Commission-sharing agreements"),
    ("Feed-ready integration architecture", "Guaranteed external portal syndication before agreements"),
], [4680, 4680])

page_break(doc)
add_heading(doc, "4. Users, hierarchy, and subscriptions", 1)
add_body(doc, "The broker or brokerage is the root of the professional tree. Public consumers and SteadFast employees exist outside the brokerage hierarchy.")
add_table(doc, ["Area", "User or role", "Relationship"], [
    ("Professional", "Broker", "Owns and controls the brokerage account; inherits broker-staff capabilities."),
    ("Professional", "Broker staff", "Receives delegated broker-like permissions and may also hold the agent role."),
    ("Professional", "Agent", "Belongs to one broker at a time and represents published listings."),
    ("Public", "Registered user", "Browses, saves, and contacts agents; not attached to a brokerage."),
    ("Public", "Unregistered visitor", "Browses and contacts agents without creating an account."),
    ("Internal", "SteadFast operations", "Handles service, billing support, system issues, and broker notifications."),
    ("Internal", "SteadFast administrator", "Controls platform configuration, staff access, integrations, and technical administration."),
], [1800, 2100, 5460])

add_heading(doc, "4.1 Multiple roles and billing", 2)
add_bullets(doc, [
    "Roles are permissions attached to one person, not separate accounts.",
    "A broker may also act as an agent.",
    "A broker staff member may also act as an agent.",
    "A person pays once even when holding multiple roles.",
    "An agent may belong to only one brokerage at a time in the first release.",
    "Broker staff with approval permission may approve their own listing submissions and changes.",
])

add_heading(doc, "4.2 Broker staff permissions", 2)
add_body(doc, "The broker may delegate permissions rather than relying on a single fixed staff role. Typical permissions include:")
add_bullets(doc, [
    "Approve new listings and material listing changes.",
    "Manage agents and account membership.",
    "Reassign unassigned listings.",
    "Review flags, approval history, and audit records.",
    "Manage brokerage profile information and reports.",
])
add_callout(doc, "Protected broker authority", "The principal broker retains control over brokerage ownership, senior staff appointment, brokerage closure, and transfer of primary broker responsibility.", fill=LIGHT_GRAY)

add_heading(doc, "5. Listing ownership and representation", 1)
add_bullets(doc, [
    "The brokerage owns every approved listing record created under it.",
    "The agent is the listing representative and operational contact, not the data owner.",
    "Every public listing must have one active agent representative.",
    "A physical property and a listing are separate records; the same property may have historical or contractually valid subsequent listings.",
    "Possible duplicate properties are detected and flagged for broker review rather than automatically merged or rejected.",
    "Listings are archived rather than permanently deleted so ownership, approvals, and history remain provable.",
])

add_heading(doc, "5.1 Agent departure", 2)
add_numbered(doc, [
    "The agent's SteadFast account and personal website remain active.",
    "The former brokerage membership becomes inactive and the agent loses access to that brokerage's private records.",
    "Listings represented by the departing agent are immediately removed from public view, agent websites, shared websites, and external feeds.",
    "Those listings remain in the former brokerage's private workspace as unassigned records.",
    "The broker assigns another active representative before republication.",
    "The former agent's profile, website configuration, account history, and qualifying shared listings remain available.",
    "If the agent joins another brokerage, the same login and website continue after the new broker approves membership; former listings do not transfer.",
])

page_break(doc)
add_heading(doc, "6. Listing workflow and approvals", 1)
add_body(doc, "Agents propose listings and changes; brokers or authorized broker staff control publication. SteadFast does not participate in brokerage approval decisions.")
add_table(doc, ["Workflow state", "Meaning", "Publicly visible"], [
    ("Draft", "Agent is preparing the listing.", "No"),
    ("Pending approval", "Submitted to broker or authorized staff.", "No for new listing; current version remains live for edits"),
    ("Changes requested", "Reviewer returned comments for correction.", "No for new listing; current version remains live for edits"),
    ("Approved / Active", "Approved version is published according to visibility.", "Yes when visibility is public"),
    ("Under offer", "Status request approved by brokerage.", "Yes, with status"),
    ("Sold / Rented", "Final status approved by brokerage.", "Removed from active search"),
    ("Withdrawn / Expired", "Broker-approved removal from active marketing.", "No"),
    ("Archived", "Retained for history and reporting.", "No"),
], [1800, 4860, 2700])

add_heading(doc, "6.1 Actions requiring approval", 2)
add_bullets(doc, [
    "Initial publication.",
    "Price changes and material property edits.",
    "Address, representative, or visibility changes.",
    "Under-offer, sold, rented, withdrawn, expired, removal, and republication actions.",
    "External publication or syndication actions where the subscription tier requires approval.",
])

add_heading(doc, "6.2 Pending change design", 2)
add_body(doc, "A proposed change must not overwrite the currently published version. The broker sees a clear before-and-after comparison and may approve, reject, or request corrections. Approved changes are applied atomically; rejected and resubmitted versions remain in the permanent history.")
add_bullets(doc, [
    "Submitter, approver, timestamps, comments, and changed fields are recorded.",
    "Self-approval by authorized broker staff is permitted and explicitly identified.",
    "Agents and brokerage reviewers can view the history relevant to their listings.",
    "An approved removal archives the listing rather than erasing it.",
])

doc.add_page_break()
add_heading(doc, "7. Visibility", 1)
add_table(doc, ["Visibility", "Audience", "Rule"], [
    ("Private", "Assigned agent and authorized brokerage users", "Used for drafts, pending work, and internal records."),
    ("Professional network", "Approved SteadFast agents and brokers", "Available for professional discovery or advertising subject to listing state."),
    ("Public", "Visitors and registered users", "Requires approval, active representative, and active status."),
], [1800, 3000, 4560])

add_heading(doc, "8. Agent-to-agent sharing", 1)
add_body(doc, "Sharing is an advertising permission and does not require broker approval. It never transfers ownership or edit authority.")
add_bullets(doc, [
    "The displaying agent may show the listing on their personal website but cannot edit it.",
    "The listing page shows both the displaying agent and the listing-owner agent; the consumer chooses whom to contact.",
    "Both agents receive an inquiry notification, while the consumer-selected agent is recorded as the primary recipient.",
    "The displaying agent may remove the listing from their own website; the owner agent is notified.",
    "The owner agent may revoke a share; the displaying agent is notified.",
    "Every approved listing change notifies all displaying agents and updates every display automatically.",
    "Sold, rented, withdrawn, expired, unassigned, or unpublished listings are removed from all shared displays automatically.",
    "Sharing, removal, revocation, and notification events are retained in the audit history.",
])

page_break(doc)
add_heading(doc, "9. Public experience", 1)
add_heading(doc, "9.1 Visitor and registered-user pages", 2)
add_bullets(doc, [
    "Homepage with simple discovery paths.",
    "Property search with filters, list view, and synchronized map view.",
    "Individual property pages with the relevant agent contacts.",
    "Agent public websites and profiles.",
    "Brokerage public websites and agent directories.",
    "Inquiry and viewing-request forms without mandatory registration.",
    "Optional registered-user favourites, saved searches, and inquiry history.",
])

add_heading(doc, "9.2 Map behaviour", 2)
add_bullets(doc, [
    "Close zoom shows individual property markers.",
    "Medium zoom groups nearby markers into numbered clusters.",
    "Parish and national zoom show area clusters and listing counts.",
    "Moving or zooming the map refreshes results inside the visible area.",
    "Exact geocoded addresses are used by default, with a future broker-controlled approximate-location option if required.",
])

add_heading(doc, "10. Agent website", 1)
add_body(doc, "Every active agent account receives a persistent website and subdomain. The site remains active when the agent changes brokerages, although former-brokerage listings are removed.")
add_bullets(doc, [
    "Agent profile, biography, photograph, service areas, and contact methods.",
    "Agent branding, social links, and website preferences.",
    "Agent's assigned approved public listings.",
    "Listings shared with the agent for advertising.",
    "Search, filtering, map display, and featured listings.",
    "Both displaying-agent and listing-owner contacts on shared listing pages.",
])

add_heading(doc, "11. Brokerage website", 1)
add_body(doc, "Every brokerage receives a branded public website showing all approved public listings it owns and all active agents it chooses to display.")
add_bullets(doc, [
    "Brokerage name, logo, description, branding, locations, and social links.",
    "All approved public listings owned by the brokerage and assigned to active representatives.",
    "Search, filters, list view, clustered map, and featured listings.",
    "Agent directory and links to individual agent websites.",
    "Property inquiries directed to agents rather than the brokerage.",
    "No shared third-party listings unless the brokerage itself owns them.",
])
add_callout(doc, "Consumer contact rule", "The brokerage provides trust and attribution, but the primary action is Contact Agent or Find an Agent. There is no Contact Brokerage action on property pages.")
contact_rule_spacer = doc.paragraphs[-1]
contact_rule_spacer.paragraph_format.line_spacing = Pt(1)

add_heading(doc, "12. Secure application workspaces", 1)
add_table(doc, ["Workspace", "Primary capabilities"], [
    ("Agent", "Dashboard, create listing, my listings, pending changes, shared listings, website, inquiries, notifications, profile, subscription."),
    ("Broker staff", "Agent capabilities when applicable, approval queue, agent management, reassignment, portfolio, flags, reports, audit history."),
    ("Broker", "All staff features plus company profile, staff permissions, agent applications, billing, integrations, brokerage website, full company reporting."),
    ("SteadFast operations", "Support, billing assistance, account help, technical issues, feed monitoring, flag notification, and response tracking."),
    ("SteadFast admin", "Internal staff permissions, platform configuration, plans, integrations, security logs, system monitoring, and administrative tools."),
], [2100, 7260])

page_break(doc)
add_heading(doc, "13. SteadFast operational boundaries", 1)
add_bullets(doc, [
    "SteadFast operations does not approve brokerage listings or material changes.",
    "SteadFast does not decide listing, ownership, commission, or agent disputes; these go to the responsible broker.",
    "SteadFast may monitor flags, notify the broker, and record whether the broker responded.",
    "SteadFast handles system issues, access support, subscription support, integration failures, and platform communications.",
    "A narrowly defined emergency legal or security content-control process remains a future policy decision.",
])

add_heading(doc, "14. Integration and API direction", 1)
add_body(doc, "SteadFast will maintain its own canonical listing model and an integration layer rather than shaping the database around one external portal.")
add_bullets(doc, [
    "Use RESO-aligned concepts and field mappings where practical.",
    "Treat RAJ/MLS and Realtor.com International/Move as governed distribution channels requiring authorization and feed agreements.",
    "Track an independent external identifier, eligibility result, validation errors, delivery state, and update history for each channel.",
    "Keep external adapters separate so future portals can be added without changing listing ownership and approval logic.",
    "Never assume technical feed compatibility grants permission to distribute listings.",
])

add_heading(doc, "15. Proposed technology architecture", 1)
add_table(doc, ["Layer", "Initial choice", "Portability approach"], [
    ("Web application", "Next.js, React, TypeScript", "Run on Vercel, Node.js, or Docker-compatible hosts."),
    ("Database", "Supabase Postgres", "Versioned SQL migrations and standard Postgres data model."),
    ("Authentication", "Supabase Auth", "Isolate provider-specific access behind server-side modules."),
    ("File storage", "Supabase Storage", "Storage adapter for future S3-compatible providers."),
    ("Geospatial", "PostGIS", "Standard Postgres extension for points, bounds, and map queries."),
    ("Deployment", "Vercel previews and staging", "Production-ready Docker/Node configuration maintained."),
    ("Edge and DNS", "Cloudflare when introduced", "DNS, CDN, WAF, and routing kept separate from business logic."),
], [1800, 3000, 4560])

add_heading(doc, "16. Environment and deployment model", 1)
add_table(doc, ["Environment", "Purpose", "Data"], [
    ("Local", "Developer implementation and tests", "Local generated sample data"),
    ("Feature preview", "Review each pushed branch or pull request", "Safe preview data; isolated branches when cost-effective"),
    ("Client demo / staging", "Stable client demonstration at a permanent address", "Dedicated test data and accounts"),
    ("Production", "Live SteadFast service", "Separate production database, storage, secrets, backups, and users"),
], [1800, 3900, 3660])
add_body(doc, "Every committed branch change should produce a preview deployment. Only changes that pass automated checks and review are promoted to the stable client demo. Production deployment will be tied to the protected main branch.")

page_break(doc)
add_heading(doc, "17. Git and change management", 1)
add_bullets(doc, [
    "Create a private GitHub repository owned by a SteadFast company organization.",
    "Use main for production-ready code, staging for the stable client demo, and named branches for features and fixes.",
    "Require pull requests and automated checks before merging protected branches.",
    "Commit the dependency lockfile and use reproducible dependency installation in automation.",
    "Never commit secrets, local environment files, production data, or credentials.",
    "Represent every database change through reviewed migration files.",
    "Make every deployment traceable to an exact Git commit.",
])

add_heading(doc, "18. Technical documentation standard", 1)
p = add_body(doc, "Documentation will live beside the code and be updated as part of feature completion. The repository should include:")
keep_with_next(p)
add_table(doc, ["Documentation area", "Contents"], [
    ("Product", "Requirements, roles, permissions, workflows, and acceptance rules."),
    ("Architecture", "System overview, data model, multi-tenancy, service boundaries, and architecture decisions."),
    ("API", "Internal endpoints, external feeds, webhooks, authentication, validation, and examples."),
    ("Operations", "Local setup, deployment, monitoring, backups, recovery, troubleshooting, and incident handling."),
    ("Security", "Security model, access control, threat model, audit approach, and incident response."),
    ("Onboarding", "Developer environment, repository workflow, test data, conventions, and first-day checklist."),
], [2400, 6960])
add_callout(doc, "Definition of done", "A feature is not complete until its tests, migrations, operational notes, and relevant product or technical documentation have been updated.", fill=LIGHT_GRAY)

add_heading(doc, "19. Security baseline", 1)
add_body(doc, "Security is part of implementation and review from the first commit. The application will follow secure-by-default TypeScript, React, Next.js, and Supabase practices.")
add_bullets(doc, [
    "Enforce authentication and authorization on the server and in database row-level security policies.",
    "Use runtime validation for every request, form, file, identifier, and external payload.",
    "Keep secret keys server-side; treat all browser-exposed configuration as public.",
    "Use explicit database grants, least-privilege policies, and audited administrative paths.",
    "Restrict uploads by type, content, size, path, and access; never serve active user content unsafely.",
    "Protect login, registration, inquiry, upload, and expensive endpoints with rate and resource limits.",
    "Use secure session handling, cross-site request protections, safe redirects, and strict origin validation.",
    "Apply browser security headers, including a practical Content Security Policy.",
    "Redact secrets and sensitive data from logs and user-facing errors.",
    "Pin dependencies, commit lockfiles, scan for vulnerable packages, and promptly apply security updates.",
    "Perform a threat model before the external pilot and a formal security audit before production launch.",
])

add_heading(doc, "20. Delivery phases", 1)
add_table(doc, ["Phase", "Outcome"], [
    ("1. Foundation", "Repository, environments, authentication, organizations, roles, permissions, audit foundation, and documentation structure."),
    ("2. Listings", "Property and listing records, media, drafts, approvals, change comparisons, status transitions, and broker ownership."),
    ("3. Public discovery", "Public search, property pages, map clusters, inquiries, agent profiles, and brokerage profiles."),
    ("4. Websites and sharing", "Agent and brokerage subdomains, sharing permissions, dual-agent contact selection, and notifications."),
    ("5. Operations", "SteadFast support, account help, billing-support placeholders, flags, feed monitoring, and administrative controls."),
    ("6. External distribution", "Authorized RAJ/MLS and Realtor.com International/Move adapters, validation, monitoring, and subscription-tier controls."),
], [2100, 7260])

page_break(doc)
add_heading(doc, "21. Confirmed decisions", 1)
add_bullets(doc, [
    "SteadFast focuses first on property listing and marketing, not tenant or rental-property operations.",
    "Vacation rentals are excluded from the initial scope; residential, long-term rental, commercial, land, and developments are included.",
    "Brokerages approve listings and all material changes; SteadFast does not participate in approval or disputes.",
    "Brokerages own listings, and each published listing requires an active agent representative.",
    "Agent accounts and websites persist when the agent changes brokerages, but former-broker listings are removed from public display.",
    "Agent-to-agent sharing requires no broker approval and grants display permission only.",
    "Consumers choose between the displaying agent and owner agent on a shared listing; both receive notifications.",
    "Brokers, broker staff, and agents may be combined on one account with one subscription payment.",
    "Broker staff may approve their own submissions when granted approval permission.",
    "Brokerage websites show only approved public listings owned by that brokerage and direct inquiries to agents.",
    "The codebase, documentation, migrations, tests, and deployment configuration are maintained in Git.",
    "Every committed feature branch should receive a client-viewable preview deployment.",
])

add_heading(doc, "22. Future decisions", 1)
add_bullets(doc, [
    "Subscription tier names, prices, grace periods, payment methods, and which distribution channels belong to each tier.",
    "Exact broker-staff permission presets and whether custom permissions are required at launch.",
    "Consumer account timing for favourites, saved searches, and notifications.",
    "Listing expiry policy and extension workflow.",
    "Email, in-app, SMS, and WhatsApp notification rollout.",
    "Emergency legal or security authority for SteadFast administrators to hide content in exceptional circumstances.",
    "Permanent SteadFast domain, custom brokerage domains, and launch branding.",
    "Formal data-feed agreements and certification requirements for RAJ/MLS and Realtor.com International/Move.",
    "Production hosting provider, backup targets, monitoring provider, and service-level objectives.",
])

add_heading(doc, "23. Initial success criteria", 1)
add_bullets(doc, [
    "A new agent can join one brokerage and create a listing without training.",
    "A broker or authorized staff member can review changes, see differences, approve or return them, and inspect the full history.",
    "No unapproved or unassigned listing can appear publicly.",
    "A public visitor can search Jamaica listings on a map and contact an agent without registering.",
    "Agent and brokerage websites automatically reflect approved listing changes and removals.",
    "Sharing preserves ownership, prevents recipient editing, presents both agents, and sends the required notifications.",
    "Professional data remains isolated between brokerages and protected by server and database authorization.",
    "Every release is traceable, deployable, documented, and portable to a standard Node.js or Docker environment.",
])

add_heading(doc, "24. Reference direction", 1)
add_bullets(doc, [
    "Realtors Association of Jamaica bylaws and MLS policies: https://realtorsjamaica.org/",
    "Real Estate Board of Jamaica registration and licensing information: https://reb.gov.jm/",
    "Real Estate Standards Organization data standards: https://www.reso.org/",
    "Realtor.com / Move listing-feed direction: https://www.move.com/ and https://b2bweb.realtor.com/",
    "Vercel Git deployment documentation: https://vercel.com/docs/git",
    "Supabase documentation and changelog: https://supabase.com/docs and https://supabase.com/changelog",
])

doc.core_properties.title = "SteadFast Product and Technical Plan"
doc.core_properties.subject = "Product requirements and technical planning baseline"
doc.core_properties.author = "SteadFast"
doc.core_properties.keywords = "SteadFast, real estate, listings, broker, agent, SaaS, Jamaica"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
